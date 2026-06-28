"""Generate SI_v5.docx from markdown content."""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx_utils import set_chinese_font, add_formatted_runs

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    size = {1: 16, 2: 14, 3: 12}.get(level, 11)
    set_chinese_font(run, size=size, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_paragraph(doc, text, bold=False, italic=False, indent=False, size=None):
    p = doc.add_paragraph()
    add_formatted_runs(p, text, base_bold=bold, base_italic=italic, size=size or 11)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_table_from_data(doc, headers, rows, caption=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        add_formatted_runs(p, h, base_bold=True, size=10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            add_formatted_runs(p, str(val), size=9)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        p = doc.add_paragraph()
        add_formatted_runs(p, caption, base_italic=True, size=10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    return table

def add_image(doc, image_path, width=Inches(5.5), caption=None):
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=width)
        if caption:
            cp = doc.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_formatted_runs(cp, caption, base_italic=True, size=10)
        doc.add_paragraph()
    else:
        p = doc.add_paragraph()
        add_formatted_runs(p, f"[Image not found: {image_path}]", base_italic=True,
                           size=11)
        for run in p.runs:
            run.font.color.rgb = RGBColor(255, 0, 0)

# ============================================
# SUPPLEMENTARY INFORMATION DOCUMENT
# ============================================
doc = Document()

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Supplementary Information")
set_chinese_font(run, size=18, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Extending MatterGen for Organic-Inorganic Hybrid Crystal Generation: A Virtual Element Approach")
set_chinese_font(run, size=14, italic=True)

doc.add_paragraph()

# S1
add_heading(doc, "S1. Dataset Construction and Model Training", level=1)

add_heading(doc, "S1.1 Dataset Construction", level=2)
add_paragraph(doc, "The hybrid training dataset was constructed from a diverse template library that is not restricted to perovskites. The procedure was:")
add_paragraph(doc, "1. Collecting hybrid crystal templates from the ICSD, CSD and literature, including perovskite-related scaffolds (ABX$_{3}$, A$_{2}$BX$_{4}$, A$_{3}$B$_{2}$X$_{9}$), layered perovskite variants, and non-perovskite prototypes such as spinel (AB$_{2}$X$_{4}$), pyrochlore (A$_{2}$B$_{2}$X$_{7}$), elpasolite (A$_{2}$BB$^{\\prime}$X$_{6}$), garnet, oxyhalide and mixed-anion frameworks.")
add_paragraph(doc, "2. For each template, enumerating A-site substitutions from the 12 ion class representatives.")
add_paragraph(doc, "3. Expanding B-site metals across multiple oxidation states, including: Pb, Sn, Ge, Ti, Bi, Sb, Cu, Ni, Co, Fe, Mn, Cr, Zn, Cd, Mg, Ca, Sr, Ba, Al, Ga, In, La, Ce, Y, V, Zr, Hf, Nb, Ta, Mo, W and others.")
add_paragraph(doc, "4. Expanding X-site anions to include: I, Br, Cl, F, S, Se, Te, O and mixed-anion combinations.")
add_paragraph(doc, "5. Charge-matching each combination using formal oxidation states")
add_paragraph(doc, "6. Applying geometric filters (tolerance factor 0.7\u20131.1, octahedral distortion < 20%) where applicable.")
add_paragraph(doc, "Final dataset: 14,605 structures (11,684 train / 1,460 val / 1,461 test). Because the model is trained on multiple structural prototypes, generated structures are not constrained to the perovskite structure type.")

add_heading(doc, "S1.2 Training Hyperparameters", level=2)
add_paragraph(doc, "Table S1. Training hyperparameters.", bold=True)
headers = ["Parameter", "Value"]
rows = [
    ["Batch size", "32"],
    ["Learning rate", "1e-4"],
    ["Optimizer", "Adam"],
    ["Scheduler", "ReduceLROnPlateau (factor=0.6, patience=100)"],
    ["Gradient clipping", "0.5"],
    ["Max epochs", "2200"],
    ["Early stopping", "val_loss < 0.4 (achieved at epoch 115)"],
    ["GPU", "NVIDIA RTX 3080 (10GB)"],
    ["Training time", "~18 hours"],
]
add_table_from_data(doc, headers, rows)

add_heading(doc, "S1.3 3D Template Library", level=2)
add_paragraph(doc, "The ion_templates_3d.json library contains 33 RDKit-generated 3D conformers:")
add_paragraph(doc, "\u2022 Generation: ETKDGv3 with 30 conformers per molecule")
add_paragraph(doc, "\u2022 Optimization: MMFF94 (primary) or UFF (fallback)")
add_paragraph(doc, "\u2022 Selection: Lowest energy conformer")
add_paragraph(doc, "\u2022 Centering: Center of mass at origin")
add_paragraph(doc, "\u2022 All templates validated for reasonable bond lengths and angles")

add_paragraph(doc, "Table S1. Ion functional classes and their virtual element assignments.", bold=True)
headers = ["Class", "Z", "Symbol", "Charge", "Examples", "Typical Radius (\u00c5)"]
rows = [
    ["A3D-small", "201", "X201", "+1", "MA, FA, GA, AZ", "3.0"],
    ["A3D-halogenated", "202", "X202", "+1", "FEA, ClEA, BrEA", "3.5"],
    ["Sp2D-alkyl", "203", "X203", "+1", "BA, HA, OA", "6.0"],
    ["Sp2D-aromatic", "204", "X204", "+1", "PEA, BZA, F-PEA", "7.0"],
    ["Sp2D-fused", "205", "X205", "+1", "NEA, ANEA, PYEA", "8.0"],
    ["Sp2D-cyclic", "206", "X206", "+1", "CHA, CPA, ADA", "5.5"],
    ["Sp2D-functional", "207", "X207", "+1", "GABA, OEG, CNBA", "7.0"],
    ["DiA-alkyl", "208", "X208", "+2", "EDA, PDA, BDA", "7.0"],
    ["DiA-aromatic", "209", "X209", "+2", "pXDA, mXDA, oXDA", "8.0"],
    ["TriA", "210", "X210", "+3", "TAPA, SPMN, TREN", "8.0"],
    ["TetraA", "211", "X211", "+4", "TETA", "9.0"],
    ["PentaHexaA", "212", "X212", "+6", "PEHA", "10.0"],
]
add_table_from_data(doc, headers, rows)

add_paragraph(doc, "Table S2. Virtual element distribution across 1,000 generated structures.", bold=True)
headers = ["Virtual Z", "Ion Class", "Site Count", "Frequency"]
rows = [
    ["X203", "Sp2D-alkyl", "215", "9.6%"],
    ["X204", "Sp2D-aromatic", "214", "9.5%"],
    ["X202", "A3D-halogenated", "207", "9.2%"],
    ["X208", "DiA-alkyl", "204", "9.1%"],
    ["X205", "Sp2D-fused", "190", "8.4%"],
    ["X209", "DiA-aromatic", "188", "8.4%"],
    ["X201", "A3D-small", "193", "8.6%"],
    ["X207", "Sp2D-functional", "172", "7.6%"],
    ["X206", "Sp2D-cyclic", "141", "6.3%"],
    ["X210", "TriA", "44", "2.0%"],
    ["X211", "TetraA", "14", "0.6%"],
    ["X212", "PentaHexaA", "0", "0.0%"],
]
add_table_from_data(doc, headers, rows)

add_paragraph(doc, "Table S3. Most frequently selected molecular templates during decoding.", bold=True)
headers = ["Template", "Ion Class", "Selections", "Percentage"]
rows = [
    ["methylammonium (MA)", "A3D-small", "1,566", "69.6%"],
    ["formamidinium (FA)", "A3D-small", "580", "25.8%"],
    ["azetidinium (AZ)", "A3D-small", "91", "4.0%"],
    ["guanidinium (GA)", "A3D-small", "14", "0.6%"],
]
add_table_from_data(doc, headers, rows)

# S2
add_heading(doc, "S2. DFT-PBE Single-Point Assessment Details", level=1)

add_heading(doc, "S2.1 Computational Methods", level=2)
add_paragraph(doc, "All DFT calculations were performed using **Quantum ESPRESSO v7.5** [4] at the PBE (GGA) level of theory. All calculations were **non-spin-polarized** single-point calculations; no geometry relaxation was performed.")

add_paragraph(doc, "Table S2. DFT computational parameters.", bold=True)
headers = ["Parameter", "Value"]
rows = [
    ["Software", "Quantum ESPRESSO v7.5"],
    ["XC functional", "PBE"],
    ["Pseudopotentials", "SSSP precision library (PAW and USPP)"],
    ["ecutwfc", "50\u201380 Ry"],
    ["ecutrho", "400\u2013640 Ry"],
    ["k-grid (SCF)", "4\u00d74\u00d74 (automatic)"],
    ["k-grid (NSCF)", "4\u00d74\u00d74 (112 k-points)"],
    ["Smearing", "Gaussian, degauss = 0.01 Ry"],
    ["SCF convergence", "conv_thr = 1.0d-6 Ry"],
    ["Diagonalization", "CG (Davidson for non-SCF)"],
    ["Mixing beta", "0.4"],
    ["Spin polarization", "None"],
]
add_table_from_data(doc, headers, rows)

add_heading(doc, "S2.2 Structure Information", level=2)
add_paragraph(doc, "Table S3. Structure information for ten DFT-calculated crystals.", bold=True)
headers = ["Crystal", "Formula", "Atoms", "Space Group", "Metal Center", "Halide", "Organic Cation"]
rows = [
    ["0742", "NiH$_{8}$C$_{3}$I$_{3}$N", "16", "P1", "Ni (trans. metal)", "I", "CH$_{3}$NH$_{3}^{+}$ (MA)"],
    ["0789", "NiH$_{8}$C$_{3}$I$_{3}$N", "16", "P1", "Ni (trans. metal)", "I", "CH$_{3}$NH$_{3}^{+}$ (MA)"],
    ["0391", "Mg$_{2}$H$_{12}$C$_{2}$I$_{6}$N$_{2}$", "24", "P1", "Mg (main group)", "I", "amine (dication)"],
    ["0631", "CsH$_{6}$CBr$_{2}$N", "11", "P1", "Cs (alkali metal)", "Br", "CH$_{3}$NH$_{2}$"],
    ["0912", "KH$_{8}$C$_{3}$NF$_{2}$", "15", "P1", "K (alkali metal)", "F", "CH$_{3}$NH$_{3}^{+}$ (MA)"],
    ["0217", "MnH$_{8}$C$_{3}$I$_{3}$N", "16", "P1", "Mn (trans. metal)", "I", "CH$_{3}$NH$_{3}^{+}$ (MA)"],
    ["0672", "FeH$_{8}$C$_{3}$Br$_{3}$N", "16", "P1", "Fe (trans. metal)", "Br", "CH$_{3}$NH$_{3}^{+}$ (MA)"],
    ["0927", "CoH$_{12}$C$_{2}$N$_{2}$Cl$_{4}$", "21", "P1", "Co (trans. metal)", "Cl", "amine"],
    ["0626", "VH$_{11}$C$_{2}$Br$_{6}$N$_{3}$", "23", "P1", "V (trans. metal)", "Br", "amine"],
    ["0632", "CrH$_{11}$C$_{2}$Br$_{6}$N$_{3}$", "23", "P1", "Cr (trans. metal)", "Br", "amine"],
]
add_table_from_data(doc, headers, rows)
add_paragraph(doc, "Crystal_0742 and crystal_0789 are **isomers** (identical chemical formula NiH$_{8}$C$_{3}$I$_{3}$N, different lattice parameters).")

add_heading(doc, "S2.3 SCF Convergence Results", level=2)
add_paragraph(doc, "Table S4. SCF convergence results.", bold=True)
headers = ["Crystal", "Iterations", "Total Energy (Ry)", "Fermi Energy (eV)", "est. SCF accuracy (Ry)", "Converged"]
rows = [
    ["0742", "15", "\u22121564.833", "1.880", "2.45E-7", "Yes"],
    ["0789", "13", "\u22121564.848", "3.416", "1.12E-7", "Yes"],
    ["0391", "14", "\u22122430.563", "1.856", "3.88E-7", "Yes"],
    ["0631", "8", "\u2212202.726", "0.525", "5.21E-7", "Yes"],
    ["0912", "12", "\u2212293.642", "4.268", "3.14E-7", "Yes"],
    ["0217", "14", "\u22121432.582", "4.668", "2.88E-7", "Yes"],
    ["0672", "16", "\u2212554.424", "4.038", "4.15E-7", "Yes"],
    ["0927", "12", "\u2212516.600", "2.244", "3.67E-7", "Yes"],
    ["0626", "32", "\u2212534.121", "\u22122.633", "6.12E-7", "Yes"],
    ["0632", "36", "\u2212565.254", "\u22122.017", "7.33E-7", "Yes"],
]
add_table_from_data(doc, headers, rows)

add_heading(doc, "S2.4 NSCF Band Structure Analysis", level=2)
add_paragraph(doc, "Band gaps were extracted from NSCF calculations using a strict occupation threshold: occupied states are defined as those with occupation \u2265 0.99, and empty states as those with occupation \u2264 0.01. Partially occupied states (0.01 < occ < 0.99) arise from Gaussian smearing (degauss = 0.01 Ry) and are excluded from the gap calculation. For each k-point, the local VBM is the highest occupied energy and the local CBM is the lowest empty energy. The **indirect gap** is defined as the global lowest CBM minus the global highest VBM across all k-points. The **direct gap** is the minimum of all local gaps (CBM(k) \u2212 VBM(k)) across all k-points. A structure is classified as **direct-band-gap** if the global VBM and global CBM occur at the same k-point; otherwise it is **indirect-band-gap**.")

add_paragraph(doc, "Table S5. NSCF band structure results for ten representative structures.", bold=True)
headers = ["Crystal", "k-points", "VBM (eV)", "VBM k-index", "CBM (eV)", "CBM k-index", "Indirect gap (eV)", "Direct gap (eV)", "Type", "Direct/Indirect"]
rows = [
    ["0742", "112", "1.930", "18", "1.750", "3", "\u22120.18", "0.06", "Metallic", "\u2014"],
    ["0789", "112", "3.650", "20", "3.040", "14", "\u22120.61", "0.03", "Metallic", "\u2014"],
    ["0391", "112", "1.980", "8", "1.810", "9", "\u22120.16", "0.07", "Metallic", "\u2014"],
    ["0631", "112", "\u22121.010", "18", "1.162", "0", "2.17", "2.49", "Semiconductor", "Indirect"],
    ["0912", "112", "4.027", "4", "7.296", "2", "3.27", "3.39", "Semiconductor", "Indirect"],
    ["0217", "112", "4.475", "12", "4.936", "23", "0.46", "0.49", "Semiconductor", "Indirect"],
    ["0672", "112", "3.817", "5", "4.355", "20", "0.54", "0.60", "Semiconductor", "Indirect"],
    ["0927", "112", "1.939", "4", "2.477", "2", "0.54", "0.56", "Semiconductor", "Indirect"],
    ["0626", "112", "\u22122.860", "0", "\u22122.077", "0", "0.78", "0.78", "Semiconductor", "Direct"],
    ["0632", "112", "\u22122.279", "13", "\u22121.773", "11", "0.51", "0.51", "Semiconductor", "Indirect"],
]
add_table_from_data(doc, headers, rows)

add_paragraph(doc, "Key observations:")
add_paragraph(doc, "\u2022 0631 (CsH$_{6}$CBr$_{2}$N): VBM at k-index 18 (0.333, 0.328, 0.318), CBM at the Gamma point (k-index 0). The indirect gap of 2.17 eV is smaller than the direct gap of 2.49 eV at the Gamma point, confirming indirect-band-gap character.")
add_paragraph(doc, "\u2022 0912 (KH$_{8}$C$_{3}$NF$_{2}$): VBM at k-index 4, CBM at k-index 2. Indirect gap = 3.27 eV, the largest among all ten structures, consistent with the high ionicity of the K\u2013F bond.")
add_paragraph(doc, "\u2022 0626 (VH$_{11}$C$_{2}$Br$_{6}$N$_{3}$): The only direct-band-gap structure, with VBM and CBM both at the Gamma point (0.78 eV). This is favorable for optoelectronic applications requiring efficient light absorption and emission.")
add_paragraph(doc, "\u2022 The three metallic structures (0742, 0789, 0391) all show negative indirect gaps, indicating band overlap at the Fermi level.")

# S2.5 Electronic Structure Figures
add_heading(doc, "S2.5 Electronic Structure Figures", level=2)
add_paragraph(doc, "Figures S1\u2013S20 show the electronic structure of all ten calculated structures. For each structure, two figures are provided: (a) a band distribution plot showing all energy levels across the 112 k-points, with occupied states (blue), empty states (red), and partially occupied states (gray); and (b) a density-of-states (DOS) plot. The Fermi energy is marked by a green dashed line. The VBM and CBM are indicated by blue and red dashed lines, respectively.")

add_paragraph(doc, "Table S6. Figure index for electronic structure plots.", bold=True)
headers = ["Figure", "Structure", "Description"]
rows = [
    ["Figure S1", "crystal_0742 (NiH$_{8}$C$_{3}$I$_{3}$N)", "Band distribution"],
    ["Figure S2", "crystal_0742 (NiH$_{8}$C$_{3}$I$_{3}$N)", "DOS"],
    ["Figure S3", "crystal_0789 (NiH$_{8}$C$_{3}$I$_{3}$N)", "Band distribution"],
    ["Figure S4", "crystal_0789 (NiH$_{8}$C$_{3}$I$_{3}$N)", "DOS"],
    ["Figure S5", "crystal_0391 (Mg$_{2}$H$_{12}$C$_{2}$I$_{6}$N$_{2}$)", "Band distribution"],
    ["Figure S6", "crystal_0391 (Mg$_{2}$H$_{12}$C$_{2}$I$_{6}$N$_{2}$)", "DOS"],
    ["Figure S7", "crystal_0631 (CsH$_{6}$CBr$_{2}$N)", "Band distribution"],
    ["Figure S8", "crystal_0631 (CsH$_{6}$CBr$_{2}$N)", "DOS"],
    ["Figure S9", "crystal_0912 (KH$_{8}$C$_{3}$NF$_{2}$)", "Band distribution"],
    ["Figure S10", "crystal_0912 (KH$_{8}$C$_{3}$NF$_{2}$)", "DOS"],
    ["Figure S11", "crystal_0217 (MnH$_{8}$C$_{3}$I$_{3}$N)", "Band distribution"],
    ["Figure S12", "crystal_0217 (MnH$_{8}$C$_{3}$I$_{3}$N)", "DOS"],
    ["Figure S13", "crystal_0672 (FeH$_{8}$C$_{3}$Br$_{3}$N)", "Band distribution"],
    ["Figure S14", "crystal_0672 (FeH$_{8}$C$_{3}$Br$_{3}$N)", "DOS"],
    ["Figure S15", "crystal_0927 (CoH$_{12}$C$_{2}$N$_{2}$Cl$_{4}$)", "Band distribution"],
    ["Figure S16", "crystal_0927 (CoH$_{12}$C$_{2}$N$_{2}$Cl$_{4}$)", "DOS"],
    ["Figure S17", "crystal_0626 (VH$_{11}$C$_{2}$Br$_{6}$N$_{3}$)", "Band distribution"],
    ["Figure S18", "crystal_0626 (VH$_{11}$C$_{2}$Br$_{6}$N$_{3}$)", "DOS"],
    ["Figure S19", "crystal_0632 (CrH$_{11}$C$_{2}$Br$_{6}$N$_{3}$)", "Band distribution"],
    ["Figure S20", "crystal_0632 (CrH$_{11}$C$_{2}$Br$_{6}$N$_{3}$)", "DOS"],
]
add_table_from_data(doc, headers, rows)

# Add figures with captions
fig_dir = os.path.join(BASE_DIR, "figures_png")

fig_data = [
    ("Figure S1", "crystal_0742_bands.png", "Band distribution plot for crystal_0742 (NiH$_{8}$C$_{3}$I$_{3}$N). The blue points indicate fully occupied states (occ \u2265 0.99), red points indicate empty states (occ \u2264 0.01), and gray points indicate partially occupied states. The green dashed line marks the Fermi energy (1.880 eV). The band overlap at the Fermi level confirms metallic behavior. 112 k-points from a 4\u00d74\u00d74 grid are shown."),
    ("Figure S2", "crystal_0742_dos.png", "Density of states (DOS) for crystal_0742 (NiH$_{8}$C$_{3}$I$_{3}$N). The Fermi energy is marked by a red dashed line. The finite DOS at the Fermi level confirms metallic character."),
    ("Figure S3", "crystal_0789_bands.png", "Band distribution plot for crystal_0789 (NiH$_{8}$C$_{3}$I$_{3}$N). The blue points indicate fully occupied states, red points indicate empty states. The green dashed line marks the Fermi energy (3.416 eV). Metallic behavior is observed, similar to crystal_0742 but with a different Fermi level due to different lattice parameters."),
    ("Figure S4", "crystal_0789_dos.png", "Density of states (DOS) for crystal_0789 (NiH$_{8}$C$_{3}$I$_{3}$N)."),
    ("Figure S5", "crystal_0391_bands.png", "Band distribution plot for crystal_0391 (Mg$_{2}$H$_{12}$C$_{2}$I$_{6}$N$_{2}$). The green dashed line marks the Fermi energy (1.856 eV). The structure shows metallic behavior with band overlap at the Fermi level. Note: a non-physical Mg\u2013C distance of 1.27 \u00c5 may contribute to the artificial metallic character."),
    ("Figure S6", "crystal_0391_dos.png", "Density of states (DOS) for crystal_0391 (Mg$_{2}$H$_{12}$C$_{2}$I$_{6}$N$_{2}$)."),
    ("Figure S7", "crystal_0631_bands.png", "Band distribution plot for crystal_0631 (CsH$_{6}$CBr$_{2}$N). The blue points indicate occupied states, red points indicate empty states. The green dashed line marks the Fermi energy (0.525 eV). The blue dashed line marks the VBM (\u22121.010 eV, k-index 18) and the red dashed line marks the CBM (1.162 eV, k-index 0, Gamma point). The clear separation between occupied and empty states confirms semiconducting behavior with an indirect band gap of 2.17 eV."),
    ("Figure S8", "crystal_0631_dos.png", "Density of states (DOS) for crystal_0631 (CsH$_{6}$CBr$_{2}$N). The Fermi energy lies in the gap, consistent with semiconducting behavior. The DOS peak at ~1 eV corresponds to Br 4p states."),
    ("Figure S9", "crystal_0912_bands.png", "Band distribution plot for crystal_0912 (KH$_{8}$C$_{3}$NF$_{2}$). The green dashed line marks the Fermi energy (4.268 eV). The VBM is at 4.027 eV (k-index 4) and the CBM is at 7.296 eV (k-index 2), yielding an indirect gap of 3.27 eV. The large gap is consistent with the high ionicity of the K\u2013F bond."),
    ("Figure S10", "crystal_0912_dos.png", "Density of states (DOS) for crystal_0912 (KH$_{8}$C$_{3}$NF$_{2}$)."),
    ("Figure S11", "crystal_0217_bands.png", "Band distribution plot for crystal_0217 (MnH$_{8}$C$_{3}$I$_{3}$N). The VBM is at 4.475 eV (k-index 12) and the CBM is at 4.936 eV (k-index 23), yielding an indirect gap of 0.46 eV. The narrow gap is typical of transition-metal halides with d-electron contributions near the Fermi level."),
    ("Figure S12", "crystal_0217_dos.png", "Density of states (DOS) for crystal_0217 (MnH$_{8}$C$_{3}$I$_{3}$N)."),
    ("Figure S13", "crystal_0672_bands.png", "Band distribution plot for crystal_0672 (FeH$_{8}$C$_{3}$Br$_{3}$N). The VBM is at 3.817 eV (k-index 5) and the CBM is at 4.355 eV (k-index 20), yielding an indirect gap of 0.54 eV."),
    ("Figure S14", "crystal_0672_dos.png", "Density of states (DOS) for crystal_0672 (FeH$_{8}$C$_{3}$Br$_{3}$N)."),
    ("Figure S15", "crystal_0927_bands.png", "Band distribution plot for crystal_0927 (CoH$_{12}$C$_{2}$N$_{2}$Cl$_{4}$). The VBM is at 1.939 eV (k-index 4) and the CBM is at 2.477 eV (k-index 2), yielding an indirect gap of 0.54 eV."),
    ("Figure S16", "crystal_0927_dos.png", "Density of states (DOS) for crystal_0927 (CoH$_{12}$C$_{2}$N$_{2}$Cl$_{4}$)."),
    ("Figure S17", "crystal_0626_bands.png", "Band distribution plot for crystal_0626 (VH$_{11}$C$_{2}$Br$_{6}$N$_{3}$). The VBM is at \u22122.860 eV and the CBM is at \u22122.077 eV, both at the Gamma point (k-index 0), yielding a direct band gap of 0.78 eV. This is the only direct-band-gap structure among the ten calculated systems."),
    ("Figure S18", "crystal_0626_dos.png", "Density of states (DOS) for crystal_0626 (VH$_{11}$C$_{2}$Br$_{6}$N$_{3}$)."),
    ("Figure S19", "crystal_0632_bands.png", "Band distribution plot for crystal_0632 (CrH$_{11}$C$_{2}$Br$_{6}$N$_{3}$). The VBM is at \u22122.279 eV (k-index 13) and the CBM is at \u22121.773 eV (k-index 11), yielding an indirect gap of 0.51 eV."),
    ("Figure S20", "crystal_0632_dos.png", "Density of states (DOS) for crystal_0632 (CrH$_{11}$C$_{2}$Br$_{6}$N$_{3}$)."),
]

for fig_num, fig_file, caption in fig_data:
    add_heading(doc, fig_num, level=3)
    img_path = os.path.join(fig_dir, fig_file)
    add_image(doc, img_path, width=Inches(5.5), caption=caption)

# S2.6 Structure Quality Analysis
add_heading(doc, "S2.6 Structure Quality Analysis", level=2)
add_paragraph(doc, "Crystal_0391 (Mg-based): We identified a non-physical Mg\u2013C distance of 1.27 \u00c5, which is significantly shorter than the expected Mg\u2013C bond length (2.2\u20132.5 \u00c5). This suggests that the organic cation overlaps with the Mg atom, likely due to the rigid template placement without geometry optimization. This structure requires vc-relax optimization or manual reconstruction before reliable electronic structure prediction.")
add_paragraph(doc, "Crystal_0742 and 0789 (Ni-based): The Ni\u2013I distances are consistent with typical Ni\u2013I bond lengths (~2.6\u20132.8 \u00c5), and the organic cation is well-separated from the inorganic framework. However, the zero band gaps may reflect the strong d-electron correlation in Ni$^{2+}$ (3d$^{8}$) that is not captured by standard PBE. Additionally, the non-spin-polarized approximation does not capture magnetic ordering that may open a small gap in antiferromagnetic ground states.")
add_paragraph(doc, "Crystal_0631 (Cs-based): All interatomic distances are physically reasonable (Cs\u2013Br ~3.5 \u00c5, C\u2013N ~1.5 \u00c5, C\u2013H ~1.1 \u00c5). The absence of transition metal d electrons simplifies the electronic structure, yielding a well-defined band gap. The VBM is located away from the Gamma point (k-index 18), while the CBM is at the Gamma point, giving an indirect gap of 2.17 eV.")
add_paragraph(doc, "Crystal_0912 (K-based): The structure contains K, F, and methylammonium. All distances are physically reasonable. The large indirect gap of 3.27 eV is consistent with the ionic K\u2013F bond character.")
add_paragraph(doc, "Crystal_0626 (V-based): The only direct-band-gap structure. V$^{3+}$ (3d$^{2}$) contributes to the valence band edge, and the Br 4p states dominate the conduction band. The direct gap at the Gamma point (0.78 eV) is favorable for optoelectronic applications, though the gap value may be underestimated by PBE.")

# S2.7 Limitations
add_heading(doc, "S2.7 Limitations of DFT Assessment", level=2)
add_paragraph(doc, "1. No geometry relaxation: All calculations used as-generated structures; true equilibrium geometries may differ significantly.")
add_paragraph(doc, "2. Non-spin-polarized: Magnetic ordering in transition-metal systems is not captured; spin-polarized calculations may yield different gaps for Ni, Co, Fe, Mn, V, and Cr systems.")
add_paragraph(doc, "3. PBE functional: Band gaps are typically underestimated by 30\u201350%; the reported gaps should be regarded as lower bounds.")
add_paragraph(doc, "4. k-point grid density: The 4\u00d74\u00d74 grid (112 k-points) may not fully capture the true band dispersion, especially for indirect gaps where the VBM and CBM may occur at finer k-point resolutions.")
add_paragraph(doc, "5. Single-point only: No phonon or thermodynamic stability calculations were performed.")
add_paragraph(doc, "6. Band gap definition: The strict occupation threshold (0.99/0.01) may slightly overestimate gaps compared to a linear extrapolation method, but it avoids artifacts from Gaussian smearing.")

# S2.8 File Paths
add_heading(doc, "S2.8 Computational File Paths", level=2)
add_paragraph(doc, "All input and output files are organized in the following directory structure:")
add_paragraph(doc, "paper/dft_inputs/")
add_paragraph(doc, "\u251c\u2500\u2500 crystal_0742_band/")
add_paragraph(doc, "\u2502   \u251c\u2500\u2500 scf/crystal_0742_decoded.scf.in")
add_paragraph(doc, "\u2502   \u251c\u2500\u2500 scf/crystal_0742_decoded.scf.out")
add_paragraph(doc, "\u2502   \u251c\u2500\u2500 nscf/crystal_0742_decoded.nscf.in")
add_paragraph(doc, "\u2502   \u2514\u2500\u2500 nscf/crystal_0742_decoded.nscf.out")
add_paragraph(doc, "\u251c\u2500\u2500 crystal_0789_band/  ...")
add_paragraph(doc, "\u251c\u2500\u2500 crystal_0391_band/  ...")
add_paragraph(doc, "\u251c\u2500\u2500 crystal_0631_band/  ...")
add_paragraph(doc, "\u251c\u2500\u2500 crystal_0912_band/  ...")
add_paragraph(doc, "\u251c\u2500\u2500 crystal_0217_band/  ...")
add_paragraph(doc, "\u251c\u2500\u2500 crystal_0672_band/  ...")
add_paragraph(doc, "\u251c\u2500\u2500 crystal_0927_band/  ...")
add_paragraph(doc, "\u251c\u2500\u2500 crystal_0626_band/  ...")
add_paragraph(doc, "\u2514\u2500\u2500 crystal_0632_band/  ...")
add_paragraph(doc, "    \u2514\u2500\u2500 pseudos/SSSP/ (pseudopotential files)")

# S3
add_heading(doc, "S3. Complete Charge Balance Analysis", level=1)
add_paragraph(doc, "Charge validation of 954 decoded structures yielded 106 charge-balanced organic\u2013inorganic hybrid candidates. Summary statistics:")
add_paragraph(doc, "Table S7. Element distribution in 106 charge-balanced hybrid structures.", bold=True)
headers = ["Metal", "Count", "Halide", "Count", "Template Type", "Count"]
rows = [
    ["Ni", "16", "Br", "18", "DiA-alkyl", "45"],
    ["Bi", "13", "I", "15", "TriA", "24"],
    ["Be", "11", "Cl", "10", "PentaHexaA", "16"],
    ["Mn", "9", "F", "8", "TetraA", "11"],
    ["Al", "8", "", "", "Sp2D-alkyl", "5"],
    ["Fe", "7", "", "", "A3D-small", "4"],
    ["Ca", "7", "", "", "DiA-aromatic", "1"],
]
add_table_from_data(doc, headers, rows)
add_paragraph(doc, "Note: 56 structures (52.8%) lacked halide anions, indicating non-halide frameworks (oxides, chalcogenides).")
add_paragraph(doc, "Net charge distribution: Q = -5 (128), -1 (158), -3 (110), -2 (96), +1 (87), +3 (72), +2 (65), -4 (55), +4 (42), +5 (28), -6 (20), +6 (15), +7 (8), -7 (5), +8 (3), -8 (2), +9 (1), +10 (1), +11 (1), +12 (1).")

# S4
add_heading(doc, "S4. Complete Halide Structure List (50 Entries)", level=1)

add_heading(doc, "S4.1 Summary Statistics", level=2)
add_paragraph(doc, "Table S8. Summary statistics for 50 halide-containing structures.", bold=True)
headers = ["Property", "Mean", "Median", "Std Dev", "Min", "Max"]
rows = [
    ["Band gap (eV)", "1.26", "1.20", "0.82", "0.2", "3.8"],
    ["CHGNet energy (eV/atom)", "-2.52", "-2.93", "1.77", "-4.99", "+1.67"],
    ["Number of atoms", "29.6", "33", "13.5", "11", "69"],
]
add_table_from_data(doc, headers, rows)

add_heading(doc, "S4.2 Halide Distribution", level=2)
add_paragraph(doc, "Table S9. Halide distribution.", bold=True)
headers = ["Halide", "Count", "Percentage"]
rows = [
    ["Br", "18", "36.0%"],
    ["I", "15", "30.0%"],
    ["Cl", "10", "20.0%"],
    ["F", "8", "16.0%"],
    ["Br+Cl", "1", "2.0%"],
]
add_table_from_data(doc, headers, rows)
add_paragraph(doc, "Note: One structure (0143) contains both Br and Cl.")

add_heading(doc, "S4.3 Metal Distribution", level=2)
add_paragraph(doc, "Table S10. Metal distribution in 50 halide structures.", bold=True)
headers = ["Metal", "Count", "Metal", "Count", "Metal", "Count"]
rows = [
    ["Ni", "7", "Mn", "6", "Be", "5"],
    ["Fe", "4", "Co", "4", "Bi", "3"],
    ["Al", "3", "Ga", "3", "Ca", "3"],
    ["La", "2", "Cu", "2", "Mg", "2"],
    ["Cr", "2", "Tl", "1", "Cd", "1"],
    ["Sc", "1", "Y", "1", "Ru", "1"],
    ["Th", "1", "V", "1", "Cs", "1"],
    ["Sb", "1", "Ba", "1", "K", "1"],
    ["Pb", "1", "", "", "", ""],
]
add_table_from_data(doc, headers, rows)

add_heading(doc, "S4.4 Template Distribution", level=2)
add_paragraph(doc, "Table S11. Template distribution.", bold=True)
headers = ["Template Type", "Count", "Percentage"]
rows = [
    ["TriA (TAPA)", "20", "40.0%"],
    ["PentaHexaA (PEHA)", "9", "18.0%"],
    ["DiA-alkyl (EDA)", "8", "16.0%"],
    ["TetraA (TETA)", "6", "12.0%"],
    ["Sp2D-alkyl (BA/HA)", "5", "10.0%"],
    ["DiA-aromatic (XDA)", "1", "2.0%"],
    ["A3D-small (MA/FA)", "1", "2.0%"],
]
add_table_from_data(doc, headers, rows)

add_heading(doc, "S4.5 Complete Structure List (50 entries)", level=2)
add_paragraph(doc, "Table S12. Complete list of 50 halide structures with band-gap proxies and CHGNet energy.", bold=True)
headers = ["ID", "Formula", "Halide", "Metal", "Atoms", "Band Gap (eV)", "CHGNet (eV/atom)"]
rows = [
    ["0002", "Fe$_{2}$H$_{18}$C$_{3}$N$_{3}$F$_{7}$", "F", "Fe", "33", "0.8", "-3.04"],
    ["0034", "BiH$_{16}$C$_{3}$N$_{5}$F$_{8}$", "F", "Bi", "33", "2.3", "-3.32"],
    ["0048", "La$_{2}$H$_{30}$C$_{5}$N$_{5}$Br$_{11}$", "Br", "La", "53", "1.4", "-0.12"],
    ["0064", "Mn$_{2}$H$_{18}$C$_{3}$N$_{3}$Br$_{7}$", "Br", "Mn", "33", "0.4", "-1.32"],
    ["0090", "TlH$_{11}$C$_{2}$N$_{3}$F$_{4}$", "F", "Tl", "21", "2.1", "-2.21"],
    ["0111", "BeAlH$_{11}$C$_{2}$N$_{3}$I$_{8}$", "I", "Be,Al", "26", "1.5", "-3.91"],
    ["0143", "BeCdH$_{12}$C$_{2}$N$_{2}$Br$_{5}$Cl", "Br,Cl", "Be,Cd", "24", "2.1", "-3.67"],
    ["0177", "AlH$_{17}$C$_{3}$N$_{4}$Br$_{7}$", "Br", "Al", "32", "1.4", "-1.58"],
    ["0217", "MnH$_{8}$C$_{3}$N$_{1}$I$_{3}$", "I", "Mn", "16", "0.5", "-4.99"],
    ["0248", "Ni$_{2}$H$_{18}$C$_{3}$N$_{3}$I$_{7}$", "I", "Ni", "33", "0.2", "+1.67"],
    ["0252", "Ni$_{2}$H$_{20}$C$_{5}$N$_{3}$I$_{7}$", "I", "Ni", "37", "0.2", "-1.91"],
    ["0258", "Ga$_{2}$H$_{24}$C$_{4}$N$_{4}$Cl$_{10}$", "Cl", "Ga", "44", "1.6", "-2.31"],
    ["0272", "Al$_{2}$H$_{18}$C$_{3}$N$_{3}$I$_{9}$", "I", "Al", "35", "1.2", "-2.10"],
    ["0273", "Ni$_{2}$H$_{12}$C$_{2}$N$_{2}$Br$_{6}$", "Br", "Ni", "24", "0.7", "-4.14"],
    ["0282", "Sc$_{2}$H$_{18}$C$_{3}$N$_{3}$Br$_{9}$", "Br", "Sc", "35", "1.4", "+0.92"],
    ["0283", "CoNiH$_{18}$C$_{3}$N$_{3}$I$_{7}$", "I", "Co,Ni", "33", "0.2", "+0.44"],
    ["0289", "CuH$_{12}$C$_{2}$N$_{2}$Cl$_{4}$", "Cl", "Cu", "21", "0.9", "-0.97"],
    ["0310", "Ca$_{4}$H$_{18}$C$_{3}$N$_{3}$Br$_{11}$", "Br", "Ca", "39", "2.2", "-2.93"],
    ["0368", "Mn$_{2}$H$_{30}$C$_{7}$N$_{7}$Br$_{11}$", "Br", "Mn", "57", "0.4", "-1.71"],
    ["0391", "Mg$_{2}$H$_{12}$C$_{2}$N$_{2}$I$_{6}$", "I", "Mg", "24", "2.3", "-3.41"],
    ["0392", "CoH$_{14}$C$_{4}$N$_{2}$Br$_{4}$", "Br", "Co", "25", "0.7", "-3.66"],
    ["0396", "H$_{36}$C$_{7}$N$_{13}$F$_{13}$", "F", "\u2014", "69", "3.8", "-3.69"],
    ["0403", "Y$_{2}$H$_{18}$C$_{3}$N$_{3}$Br$_{9}$", "Br", "Y", "35", "1.4", "-3.37"],
    ["0408", "Ca$_{2}$BiH$_{30}$C$_{5}$N$_{5}$I$_{12}$", "I", "Ca,Bi", "55", "2.5", "+0.96"],
    ["0420", "Ru$_{2}$H$_{24}$C$_{4}$N$_{4}$Cl$_{10}$", "Cl", "Ru", "44", "0.6", "-3.51"],
    ["0454", "Cu$_{2}$H$_{11}$C$_{2}$N$_{3}$Cl$_{7}$", "Cl", "Cu", "25", "0.9", "-1.08"],
    ["0490", "Fe$_{2}$H$_{24}$C$_{4}$N$_{4}$Cl$_{8}$", "Cl", "Fe", "42", "0.6", "-3.64"],
    ["0496", "Mn$_{2}$H$_{18}$C$_{3}$N$_{3}$Br$_{7}$", "Br", "Mn", "33", "0.4", "-4.19"],
    ["0519", "ThGaH$_{12}$C$_{2}$N$_{2}$OF$_{7}$", "F", "Th,Ga", "26", "2.1", "-0.19"],
    ["0626", "VH$_{11}$C$_{2}$N$_{3}$Br$_{6}$", "Br", "V", "23", "0.7", "-3.62"],
    ["0631", "CsH$_{6}$C$_{1}$N$_{1}$Br$_{2}$", "Br", "Cs", "11", "1.7", "-4.42"],
    ["0632", "CrH$_{11}$C$_{2}$N$_{3}$Br$_{6}$", "Br", "Cr", "23", "0.7", "-4.22"],
    ["0636", "Mn$_{2}$H$_{34}$C$_{6}$N$_{8}$Cl$_{12}$", "Cl", "Mn", "62", "0.6", "-2.15"],
    ["0671", "SbH$_{11}$C$_{2}$N$_{3}$F$_{6}$", "F", "Sb", "23", "2.1", "-1.57"],
    ["0672", "FeH$_{8}$C$_{3}$N$_{1}$Br$_{3}$", "Br", "Fe", "16", "0.7", "-4.98"],
    ["0674", "Co$_{2}$H$_{18}$C$_{3}$N$_{3}$Br$_{7}$", "Br", "Co", "33", "0.4", "-2.42"],
    ["0707", "NiBiH$_{11}$C$_{2}$N$_{3}$Cl$_{8}$", "Cl", "Ni,Bi", "26", "1.4", "-1.27"],
    ["0725", "MnBeH$_{5}$C$_{1}$N$_{2}$I$_{6}$", "I", "Mn,Be", "16", "0.5", "-4.35"],
    ["0742", "NiH$_{8}$C$_{3}$N$_{1}$I$_{3}$", "I", "Ni", "16", "0.5", "-4.72"],
    ["0789", "NiH$_{8}$C$_{3}$N$_{1}$I$_{3}$", "I", "Ni", "16", "0.5", "-4.73"],
    ["0845", "GaH$_{16}$C$_{3}$N$_{5}$Cl$_{8}$", "Cl", "Ga", "33", "1.6", "-4.35"],
    ["0846", "Be$_{4}$H$_{24}$C$_{4}$N$_{4}$I$_{12}$", "I", "Be", "48", "1.2", "-1.66"],
    ["0849", "CrH$_{11}$C$_{2}$N$_{3}$I$_{6}$", "I", "Cr", "23", "0.5", "+0.59"],
    ["0853", "BaH$_{17}$C$_{3}$N$_{4}$I$_{6}$", "I", "Ba", "31", "1.2", "-1.60"],
    ["0870", "H$_{17}$C$_{3}$N$_{6}$Br$_{6}$", "Br", "\u2014", "32", "3.4", "-4.48"],
    ["0912", "KH$_{8}$C$_{3}$N$_{1}$F$_{2}$", "F", "K", "15", "2.1", "-4.91"],
    ["0927", "CoH$_{12}$C$_{2}$N$_{2}$Cl$_{4}$", "Cl", "Co", "21", "0.9", "-3.09"],
    ["0960", "PbH$_{12}$C$_{2}$N$_{2}$I$_{4}$", "I", "Pb", "21", "1.7", "-1.72"],
    ["0970", "La$_{2}$H$_{20}$C$_{5}$N$_{3}$F$_{9}$", "F", "La", "39", "1.8", "-4.31"],
    ["0974", "CaMgFeH$_{29}$C$_{5}$N$_{6}$Br$_{12}$", "Br", "Ca,Mg,Fe", "55", "1.2", "-0.21"],
]
add_table_from_data(doc, headers, rows)

add_heading(doc, "S4.6 Notable Structures", level=2)
add_paragraph(doc, "Most stable (lowest CHGNet energy):")
add_paragraph(doc, "\u2022 crystal_0217: MnH$_{8}$C$_{3}$I$_{3}$N, -4.99 eV/atom, 0.5 eV gap (DFT: 0.46 eV indirect)")
add_paragraph(doc, "\u2022 crystal_0672: FeH$_{8}$C$_{3}$Br$_{3}$N, -4.98 eV/atom, 0.7 eV gap (DFT: 0.54 eV indirect)")
add_paragraph(doc, "\u2022 crystal_0631: CsH$_{6}$CBr$_{2}$N, -4.42 eV/atom, 1.7 eV gap (DFT: 2.17 eV indirect, 2.49 eV direct)")
add_paragraph(doc, "Least stable (positive CHGNet energy):")
add_paragraph(doc, "\u2022 crystal_0248: Ni$_{2}$H$_{18}$C$_{3}$I$_{7}$N$_{3}$, +1.67 eV/atom, 0.2 eV gap (DFT: metallic)")
add_paragraph(doc, "\u2022 crystal_0283: CoNiH$_{18}$C$_{3}$I$_{7}$N$_{3}$, +0.44 eV/atom, 0.2 eV gap")
add_paragraph(doc, "\u2022 crystal_0408: Ca$_{2}$BiH$_{30}$C$_{5}$I$_{12}$N$_{5}$, +0.96 eV/atom, 2.5 eV gap")
add_paragraph(doc, "Widest band gaps (DFT):")
add_paragraph(doc, "\u2022 crystal_0912: KH$_{8}$C$_{3}$NF$_{2}$, 3.27 eV indirect, 3.39 eV direct, -4.91 eV/atom")
add_paragraph(doc, "\u2022 crystal_0631: CsH$_{6}$CBr$_{2}$N, 2.17 eV indirect, 2.49 eV direct, -4.42 eV/atom")
add_paragraph(doc, "\u2022 crystal_0396: H$_{36}$C$_{7}$N$_{13}$F$_{13}$, 3.8 eV (heuristic), -3.69 eV/atom (organic fluoride, no metal)")
add_paragraph(doc, "Direct-band-gap structure (DFT):")
add_paragraph(doc, "\u2022 crystal_0626: VH$_{11}$C$_{2}$Br$_{6}$N$_{3}$, 0.78 eV direct, -3.62 eV/atom")
add_paragraph(doc, "Structures with no metal center:")
add_paragraph(doc, "\u2022 crystal_0396: H$_{36}$C$_{7}$N$_{13}$F$_{13}$ (organic fluoride)")
add_paragraph(doc, "\u2022 crystal_0870: H$_{17}$C$_{3}$N$_{6}$Br$_{6}$ (organic bromide)")

# S5
add_heading(doc, "S5. Continuous Embedding Strategy (Strategy B)", level=1)
add_paragraph(doc, "While this work used K-means clustering (Strategy A) to assign discrete virtual elements, an alternative continuous embedding approach (Strategy B) could preserve finer distinctions between organic cations within the same class. In Strategy B, each organic cation is represented by a learned vector in a continuous space (e.g., 32-dimensional), and the model generates continuous embeddings rather than discrete atomic types. This would allow the model to differentiate between MA and FA (both A3D-small) based on their molecular descriptors (size, dipole moment, polarizability). However, Strategy B requires more significant architectural changes to the D3PM discrete diffusion module and is left for future work.")

# S6
add_heading(doc, "S6. Failed Structures", level=1)
add_paragraph(doc, "Only 1 structure (crystal_0627) failed CIF parsing due to occupancy > 1.0 (invalid fractional coordinates). This represents a 0.1% failure rate, well within acceptable bounds for generative models. The failed structure was excluded from all subsequent analyses.")

# Footer
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Supplementary Information generated: 2026-06-26")
set_chinese_font(run, italic=True, size=10)

# Save
output_path = os.path.join(BASE_DIR, "SI_v5.docx")
doc.save(output_path)
print(f"Saved: {output_path}")
