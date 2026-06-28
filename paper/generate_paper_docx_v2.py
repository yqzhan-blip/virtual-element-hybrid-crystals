"""Regenerate paper_draft_v5.docx with Figure 1-5 inserted."""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx_utils import set_chinese_font, add_formatted_runs

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    size = {1: 16, 2: 14, 3: 12}.get(level, 11)
    set_chinese_font(run, size=size, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_paragraph(doc, text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    add_formatted_runs(p, text, base_bold=bold, base_italic=italic, size=11)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_table_from_data(doc, headers, rows, caption=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        add_formatted_runs(p, h, base_bold=True, size=10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            add_formatted_runs(p, str(val), size=9)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        p = doc.add_paragraph()
        add_formatted_runs(p, caption, base_italic=True, size=10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    return table

def add_image(doc, image_path, width=Inches(5.5), caption=None):
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=width)
        if caption:
            cp = doc.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_formatted_runs(cp, caption, base_italic=True, size=10)
        doc.add_paragraph()
    else:
        p = doc.add_paragraph()
        add_formatted_runs(p, f"[Image not found: {image_path}]", base_italic=True,
                           size=11)
        for run in p.runs:
            run.font.color.rgb = RGBColor(255, 0, 0)

def insert_image_after_paragraph(doc, paragraph, image_path, caption=None, width=Inches(5.5)):
    """Insert image after a specific paragraph using XML manipulation."""
    if not os.path.exists(image_path):
        print(f"  [SKIP] Not found: {os.path.basename(image_path)}")
        return
    
    # Create new paragraph for image
    p_img = doc.add_paragraph()
    p_img._element.getparent().remove(p_img._element)
    paragraph._element.addnext(p_img._element)
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(image_path, width=width)
    
    if caption:
        p_cap = doc.add_paragraph()
        p_cap._element.getparent().remove(p_cap._element)
        p_img._element.addnext(p_cap._element)
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = p_cap.add_run(caption)
        set_chinese_font(cr, italic=True, size=10)
    
    # Add spacing
    p_sp = doc.add_paragraph()
    p_sp._element.getparent().remove(p_sp._element)
    (p_cap if caption else p_img)._element.addnext(p_sp._element)
    
    print(f"  [OK] Inserted: {os.path.basename(image_path)}")

# ============================================
# BUILD DOCUMENT
# ============================================
doc = Document()

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("A virtual-element framework for generative design of organic\u2013inorganic hybrid crystals")
set_chinese_font(run, size=18, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_formatted_runs(p, "Yiqiang Zhan$^{1,2*}$", size=12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_formatted_runs(p, "$^{1}$ The State Key Laboratory of Photovoltaic Science and Technology, College of Future Information Technology, Fudan University, Shanghai 200438, China", size=10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_formatted_runs(p, "$^{2}$ Institute of Optoelectronics, Fudan University, Shanghai 200438, China", size=10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_formatted_runs(p, "$^{*}$ Corresponding author: yqzhan@fudan.edu.cn", size=10)

doc.add_paragraph()

# Abstract
add_heading(doc, "Abstract", level=1)
abstract = "Generative models trained on inorganic crystals are difficult to apply directly to organic\u2013inorganic hybrid materials because molecular organic cations are not represented in the fixed elemental vocabulary. Here we introduce a virtual-element strategy that maps organic cation families into coarse-grained pseudo-elements, enabling the fine-tuning of an inorganic diffusion model for hybrid crystal generation. We define twelve virtual ion classes based on crystal-chemical descriptors (charge, dimensionality, size, and chemical character), expanding the atom-type index space of the pretrained MatterGen model to 401 slots (including a masked gap region and 12 active virtual-element classes) while retaining pretrained inorganic embeddings. The fine-tuned model generated 1,000 virtual-element crystals, of which 954 were decoded into full-atom structures using RDKit-derived molecular templates. Formal charge analysis identified 106 charge-neutral hybrid candidates (10.6% of generated, 11.1% of decoded), including 50 halide-containing structures. High-throughput screening using machine-learning interatomic potentials and DFT-PBE single-point calculations on ten representative candidates reveal diverse electronic behavior, including a Cs-containing bromide with a 2.17 eV indirect band gap and a K\u2013F compound with a 3.27 eV indirect band gap. Only one structure (VH$_{11}$C$_{2}$Br$_{6}$N$_{3}$) exhibits a direct band gap (0.78 eV), while three Ni- and Mg-based iodides are metallic. This work establishes a transferable representation, generation and decoding framework for extending pretrained inorganic crystal diffusion models to organic\u2013inorganic hybrid material spaces, while identifying charge neutrality and structural relaxation as key bottlenecks for future development."
add_paragraph(doc, abstract, italic=True)
add_paragraph(doc, "Keywords: generative AI, organic\u2013inorganic hybrid materials, hybrid perovskites, diffusion models, virtual elements, crystal structure prediction, DFT validation", bold=True)

# Section 1
add_heading(doc, "1. Introduction", level=1)
add_paragraph(doc, "Organic\u2013inorganic hybrid crystals have emerged as one of the most promising classes of materials for optoelectronic applications, including solar cells, LEDs, and photodetectors [1,2]. The prototypical hybrid perovskite formula ABX$_{3}$ (A = organic cation, B = metal, X = halide) is among the best-known examples and offers extraordinary compositional flexibility; however, the combinatorial space of possible organic cations\u2014encompassing thousands of ammonium, phosphonium, and sulfonium species\u2014makes exhaustive experimental screening impractical across the broader hybrid crystal family. Existing generative models for crystal structure prediction, such as diffusion models trained on the Materials Project or Alexandria databases, are fundamentally limited to the fixed atom-type vocabulary of the pretrained MatterGen implementation [3]. This representational gap excludes molecular organic cations and prevents direct application of pretrained inorganic models to hybrid materials without either retraining from scratch or introducing a strategy for representing molecular components within an atomic-type framework.", indent=True)
add_paragraph(doc, "Recent advances in machine learning have accelerated both the screening and inverse design of hybrid organic\u2013inorganic materials. Supervised learning models have been developed to predict the stability, band gaps, and synthesizability of hybrid perovskites from compositional and structural descriptors [11\u201314]. Concurrently, generative models\u2014including generative adversarial networks and deep-learning-based composition predictors\u2014have begun to explore the vast chemical space of organic\u2013inorganic hybrids [15,16]. Fingerprinting strategies for organic cations have further enabled target-driven inverse design of two-dimensional hybrid perovskites [17]. These studies demonstrate that machine learning can navigate compositional and structural choices far more efficiently than exhaustive experiments. However, most existing approaches either predict properties for known compositions, generate simplified descriptors, or enumerate compositions within fixed structural templates, rather than producing full three-dimensional crystal structures with molecular organic cations embedded in an inorganic framework.", indent=True)
add_paragraph(doc, "To bridge this gap, we propose a **virtual-element** approach that compresses organic ion functional classes into extended atomic numbers, analogous to coarse-grained force fields in molecular dynamics where atom groups are represented by single interaction sites. This strategy allows us to: (i) leverage pretrained inorganic generative models without retraining from scratch, (ii) represent diverse organic cations through a compact set of ion classes, and (iii) decode generated structures back to full-atom organic\u2013inorganic hybrids using molecular templates. Here, we make three contributions. First, we introduce a virtual-element representation that maps organic cation families to coarse-grained pseudo-atomic tokens compatible with atom-type diffusion models. Second, we adapt a pretrained inorganic crystal generator using trainable virtual-element embeddings and adapter layers, avoiding training a hybrid generator from scratch. Third, we develop a decoding and filtering workflow that converts virtual-element crystals into full-atom organic\u2013inorganic structures and evaluates charge neutrality, halide chemistry, preliminary energy ranking and DFT-PBE electronic structure. Together, these components establish a transferable route for extending inorganic generative models to hybrid crystal spaces. The central contribution of this work is not merely the generation of individual candidate structures, but the introduction of a representation layer that makes molecular organic ions compatible with atom-type crystal diffusion models.", indent=True)

# Section 2
add_heading(doc, "2. Methods", level=1)

add_heading(doc, "2.1 Virtual Element Framework", level=2)
add_paragraph(doc, "We define a **virtual element** as a coarse-grained pseudo-atomic token that represents a family of organic ions with similar charge, size, shape, and crystal-chemical role. Standard chemical elements correspond to atomic numbers Z = 1\u2013118. Our initial conceptual design allocated a gap region (Z = 119\u2013200) for masking and an active virtual space (Z = 201\u2013212) for organic cations, yielding 12 ion functional classes (Table 1). These classes are defined by four crystal-chemical descriptors: (i) **charge** (monocationic, dicationic, tricationic, etc.), (ii) **dimensionality** (3D-spherical, 2D-planar, 1D-rod-like), (iii) **size** (small: radius < 4 \u00c5, medium: 4\u20137 \u00c5, large: > 7 \u00c5), and (iv) **chemical character** (alkyl, aromatic, halogenated, cyclic, functionalized). This classification maps chemically and structurally similar organic cations to the same virtual element, enabling efficient representation while preserving the structural diversity essential to hybrid materials.")
add_paragraph(doc, "Full class definitions and representative cations are provided in SI Table S2.")

add_paragraph(doc, "Table 1. The 12 ion functional classes defined for coarse-grained virtual element representation.", bold=True)
headers = ["Class", "Z", "Charge", "Dimensionality", "Size", "Examples"]
rows = [
    ["A3D-small", "201", "+1", "3D-spherical", "Small", "MA, FA, GA, AZ"],
    ["A3D-halogenated", "202", "+1", "3D-spherical", "Small", "FEA, ClEA, BrEA"],
    ["Sp2D-alkyl", "203", "+1", "2D-planar", "Medium", "BA, HA, OA"],
    ["Sp2D-aromatic", "204", "+1", "2D-planar", "Large", "PEA, BZA, F-PEA"],
    ["Sp2D-fused", "205", "+1", "2D-planar", "Large", "NEA, ANEA, PYEA"],
    ["Sp2D-cyclic", "206", "+1", "2D-planar", "Medium", "CHA, CPA, ADA"],
    ["Sp2D-functional", "207", "+1", "2D-planar", "Large", "GABA, OEG, CNBA"],
    ["DiA-alkyl", "208", "+2", "1D-rod", "Medium", "EDA, PDA, BDA"],
    ["DiA-aromatic", "209", "+2", "1D-rod", "Large", "pXDA, mXDA, oXDA"],
    ["TriA", "210", "+3", "1D-rod", "Large", "TAPA, SPMN, TREN"],
    ["TetraA", "211", "+4", "1D-rod", "Large", "TETA"],
    ["PentaHexaA", "212", "+6", "1D-rod", "Large", "PEHA"],
]
add_table_from_data(doc, headers, rows)

add_heading(doc, "2.2 Model Architecture Modifications", level=2)
add_paragraph(doc, "The MatterGen diffusion model [3] operates on a discrete vocabulary of atom types. To extend this framework to hybrid materials, we expanded the atom-type index space to 401 slots: the 101 atom types present in the pretrained model, a masked gap region (Z = 119\u2013200), and 12 active virtual element classes (Z = 201\u2013212). The 401-slot index space is reserved for compatibility with the virtual-element mapping; only the 12 virtual classes are active and trainable in the present work.")
add_paragraph(doc, "The ExtendedAtomEmbedding layer initializes the 12 virtual element embeddings by interpolating from the pretrained embeddings of chemically similar real elements (e.g., X201 from Cs and Rb for monocationic spherical ions), while freezing all pretrained real-element weights. The D3PM diffusion module samples from this extended index space, with the gap region masked so that only real elements (Z = 1\u2013118) and active virtual elements (Z = 201\u2013212) can be generated. This expansion retains the pretrained inorganic knowledge while adding only a small number of trainable virtual-element parameters. Full architectural details are in SI Section S1.")

add_heading(doc, "2.3 Training Data and Fine-tuning", level=2)
add_paragraph(doc, "We constructed a synthetic training dataset of 14,605 hybrid crystal structures by sampling from a diverse library of hybrid crystal templates. The templates include perovskite-related scaffolds (ABX$_{3}$, A$_{2}$BX$_{4}$, A$_{3}$B$_{2}$X$_{9}$) as well as non-perovskite prototypes such as spinel (AB$_{2}$X$_{4}$), pyrochlore (A$_{2}$B$_{2}$X$_{7}$), elpasolite (A$_{2}$BB$^{\\prime}$X$_{6}$), garnet, oxyhalide and mixed-anion frameworks. For each template we enumerated A-site organic cations from the 12 ion classes, B-site metals spanning multiple oxidation states, and X-site anions including halides, oxides and sulfides. The dataset was split into 11,684 training, 1,460 validation, and 1,461 test structures. To avoid overly optimistic validation arising from template leakage, structures sharing the same inorganic scaffold and organic class were assigned to the same split.")
add_paragraph(doc, "Fine-tuning used adapter layers with the base GemNet-T denoiser frozen, training only the ExtendedAtomEmbedding, adapter layers, and property embeddings. Training converged at epoch 115 with val_loss = 0.38 (SI Table S1).")

add_heading(doc, "2.4 Generation and Decoding Pipeline", level=2)
pipeline_para = add_paragraph(doc, "**Generation**: The fine-tuned model generates structures via D3PM diffusion with element masking, outputting CIF files with virtual elements (X201\u2013X212).")
add_paragraph(doc, "**Decoding**: For each virtual site, we estimate cavity radius from nearest inorganic neighbors, select the best-fitting 3D molecular template from a precomputed RDKit library (ETKDGv3 + MMFF94/UFF), rotate and center the template at the virtual site, and output a full-atom CIF. Full computational details are in SI Section S1.")

# ===== FIGURE 1: Pipeline (original style) =====
add_image(doc, 
    os.path.join(BASE_DIR, "figures", "figure_1_framework.png"),
    width=Inches(5.5),
    caption="Figure 1. Overview of the virtual-element generation and decoding pipeline. (a) Training: the pretrained MatterGen model is fine-tuned on a hybrid-structure dataset using an extended atom-type index space (401 slots, with 12 active virtual-element classes X201–X212) and adapter layers. (b) Generation: the fine-tuned model produces coarse-grained structures containing virtual elements (X201–X212). (c) Validation: virtual sites are decoded to full-atom organic cations using a precomputed 3D molecular template library, followed by charge-balance filtering and screening. The 401-slot index space includes reserved masked regions; only X201–X212 are active virtual classes in this study.")

add_heading(doc, "2.5 Formal Charge Assignment and Charge-Neutrality Filtering", level=2)
add_paragraph(doc, "Charge neutrality was evaluated by enumerating chemically plausible oxidation states for inorganic elements and fixed formal charges for organic virtual ions. The charge assignment rules are as follows:")
add_paragraph(doc, "\u2022 Organic virtual ions: fixed charges based on class (X201\u2013X207: +1; X208\u2013X209: +2; X210: +3; X211: +4; X212: +6)")
add_paragraph(doc, "\u2022 Halides: F, Cl, Br, I \u2192 \u22121")
add_paragraph(doc, "\u2022 Chalcogenides: O, S, Se, Te \u2192 \u22122")
add_paragraph(doc, "\u2022 Metals: common oxidation states from a predefined table (e.g., K: +1; Cs: +1; Ni: +2; Mn: +2/+3; Fe: +2/+3; Mg: +2; Ca: +2; Al: +3; etc.). For multi-valent metals, we enumerate all plausible combinations and check if any combination satisfies global electroneutrality.")
add_paragraph(doc, "A structure was classified as charge-neutral if at least one combination of oxidation states satisfied global electroneutrality. This formal-charge approach does not account for coordination environment or partial charge transfer, but provides a fast and chemically motivated filter. Full details are in SI Section S3.")

add_heading(doc, "2.6 DFT Single-Point Assessment", level=2)
add_paragraph(doc, "Ten representative structures were selected for DFT single-point calculations using Quantum ESPRESSO v7.5 [4] at the PBE level (SI Section S2). Structures were chosen to cover diverse chemistries: Ni-based iodides (crystal_0742 and its isomer 0789), Mg-based iodide (crystal_0391), Cs-based bromide (crystal_0631), K\u2013F compound (crystal_0912), Mn-based iodide (crystal_0217), Fe-based bromide (crystal_0672), Co-based chloride (crystal_0927), V-based bromide (crystal_0626), and Cr-based bromide (crystal_0632). All calculations were **non-spin-polarized** single-point SCF calculations followed by NSCF band structure evaluations on a 4\u00d74\u00d74 k-point grid. No geometry relaxation was performed on the generated structures; therefore, the reported band gaps represent the electronic properties of the as-generated geometries and should be interpreted as preliminary assessments rather than fully converged values. In particular, transition-metal systems (Ni, Mn, Fe, Co, V, Cr) may exhibit magnetic ordering that is not captured by non-spin-polarized calculations, and the true equilibrium geometries may differ from the as-generated structures.")

add_heading(doc, "2.7 High-Throughput Screening", level=2)
add_paragraph(doc, "To evaluate the relative energetic trends of the generated halide structures at scale, we employed CHGNet [5] for single-point energy prediction and a composition-based band-gap proxy for preliminary electronic-property ranking (SI Section S4). CHGNet provides rapid energy evaluation without explicit DFT, enabling ranking of the full 50-structure halide subset. The composition-based band-gap proxy estimates gaps based on the presence of transition metals (narrow gap), main group metals (moderate gap), and alkali/alkaline earth metals (wide gap), with corrections for halide type (I < Br < Cl < F in ionicity). While approximate, these values should be interpreted as qualitative trends for prioritizing structures for detailed DFT validation, not as accurate thermodynamic or electronic-property predictions. No clear correlation is expected between CHGNet energy and band-gap proxy, as they capture different physical properties.")

# Section 3
add_heading(doc, "3. Results", level=1)

add_heading(doc, "3.1 Generation and Decoding Statistics", level=2)
add_paragraph(doc, "We generated 1,000 hybrid crystal structures in 20 batches of 50. Of these, 954 (95.4%) were successfully decoded to full-atom representations. The 46 structures without virtual elements were pure inorganic byproducts. All 2,251 virtual sites were mapped to 3D molecular templates with 100% success rate. The generation and decoding statistics are summarized in Table 2.")

add_paragraph(doc, "Table 2. Generation and filtering pipeline statistics.", bold=True)
headers = ["Stage", "Criterion", "Count", "Fraction of generated (%)", "Fraction of previous stage"]
rows = [
    ["Generated samples", "D3PM samples", "1,000", "100.0", "\u2014"],
    ["Decoded structures", "Full-atom CIF generated", "954", "95.4", "95.4% of generated"],
    ["Charge-neutral hybrids", "Formal charge balanced", "106", "10.6", "11.1% of decoded"],
    ["Halide-containing hybrids", "Containing F/Cl/Br/I", "50", "5.0", "47.2% of charge-neutral"],
]
add_table_from_data(doc, headers, rows)
add_paragraph(doc, "Across all generated virtual sites, monocationic 2D spacer-like classes such as X203 (Sp2D-alkyl) and X204 (Sp2D-aromatic) are most frequently sampled (Fig. 2d and SI Table S3). In contrast, within the halide-containing charge-neutral subset, higher-charge rod-like templates such as TriA (X210), PentaHexaA (X212) and DiA-alkyl (X208) become enriched (Fig. 3c and SI Table S11), reflecting the stronger charge-compensation requirement of halide-rich frameworks.")

add_heading(doc, "3.2 Charge Balance Analysis", level=2)
add_paragraph(doc, "Charge validation of 954 decoded structures yielded 106 organic\u2013inorganic hybrid structures that satisfy charge neutrality (11.1% of decoded). Although only 11.1% of decoded structures satisfy formal charge neutrality, this fraction is nontrivial for an unconstrained generative model operating over both inorganic atoms and coarse-grained organic ions. The result highlights global charge conservation as a key missing inductive bias in current crystal diffusion models.")
add_paragraph(doc, "The most common net charges in unbalanced structures were \u22125 (128), \u22121 (158), and \u22123 (110), indicating negative net charges that reflect insufficient cationic charge compensation or excess anionic components in many generated structures. Notably, 56 balanced structures (52.8% of the 106) lacked halide anions, suggesting the model sometimes generates non-halide frameworks (oxides, chalcogenides) that still satisfy charge balance with organic cations. After excluding 10 pure inorganic structures, all remaining candidates contain organic cations from the 7 recognized ion classes. Full details are in SI Section S3.")

# ===== FIGURE 2: Statistics =====
add_image(doc, 
    os.path.join(BASE_DIR, "figures", "figure_2_statistics.png"),
    width=Inches(5.5),
    caption="Fig. 2 | Generation statistics and charge-neutrality filtering. a, Generation and filtering pipeline funnel showing 1,000 D3PM samples, 954 decoded structures, 106 charge-neutral hybrids and 50 halide-containing candidates (47.2% of charge-neutral hybrids). b, Net charge distribution of 954 decoded structures; 106 (11.1%) satisfy charge neutrality (Q = 0). c, Top five imbalance modes among unbalanced structures, dominated by negative net charges indicating insufficient cationic charge compensation or excess anionic components. d, Virtual element usage distribution across 2,251 virtual sites, with X203 (Sp2D-alkyl) and X204 (Sp2D-aromatic) being the most frequent.")

add_heading(doc, "3.3 Halide-containing Structures: Preliminary Energy and Band-Gap Proxy Screening", level=2)
add_paragraph(doc, "Among the 106 charge-neutral hybrid structures, 50 (47.2%) contain halide anions (F, Cl, Br, I), forming the subset most relevant to hybrid halide materials, including perovskite-inspired frameworks and related optoelectronic materials. We performed high-throughput screening on this subset using CHGNet for preliminary single-point energy ranking and a composition-based band-gap proxy for electronic-property screening (SI Section S4).")
add_paragraph(doc, "**Halide distribution.** Counting halide occurrences, bromide appears in 18 candidates, iodide in 15, chloride in 10 and fluoride in 8; one structure contains both Br and Cl. Because of this mixed-halide structure, the occurrence-level percentages (36.0%, 30.0%, 20.0% and 16.0%) sum to slightly above 100%; all 50 charge-neutral candidates contain at least one halide. The metal centers span 24 elements, with Ni (7) being most frequent, followed by Mn (6), Be (5), Fe (4) and Co (4). Notably, crystal_0631 is the sole Cs-based structure in this subset, and only one structure contains Pb (crystal_0960), highlighting the model's bias toward transition metals over the heavy p-block elements typical of experimental HOIPs.")
add_paragraph(doc, "**Metal\u2013halide and organic\u2013halide co-occurrence.** The metal\u2013halide co-occurrence map reveals that transition-metal iodides and bromides dominate the generated halide subset, whereas F-containing structures are more frequently associated with main-group or high-valent metal species. The organic-class\u2013halide map further shows enrichment of higher-charge rod-like cations in Br- and I-containing structures, consistent with their larger anion content and charge-compensation requirements.")
add_paragraph(doc, "**Band-gap proxy estimates.** The composition-based model predicts a mean band-gap proxy of 1.26 \u00b1 0.82 eV across the 50 halide structures, with a broad distribution ranging from 0.2 eV (Ni-based, crystal_0248) to 3.8 eV (organic fluoride, crystal_0396). Approximately 35% of the halide structures are predicted metallic or narrow-gap (<0.5 eV), 33% have moderate gaps (0.5\u20131.5 eV), and 31% exhibit wider gaps (>1.5 eV). The DFT-PBE single-point indirect gap of crystal_0631 (2.17 eV) falls in the upper tail of this distribution, consistent with its Cs-based, transition-metal-free composition.")
add_paragraph(doc, "**Preliminary energy ranking.** CHGNet single-point energies reveal significant variation in relative energy, ranging from \u22124.99 to +1.67 eV/atom (mean = \u22122.52 \u00b1 1.77 eV/atom). The structures with lowest predicted energies include Mn-based iodide (crystal_0217, \u22124.99 eV/atom) and Fe-based bromide (crystal_0672, \u22124.98 eV/atom), both featuring simple A-site spacers. Conversely, several Ni-based iodides (crystal_0248, +1.67 eV/atom; crystal_0283, +0.44 eV/atom) and a La-based bromide (crystal_0048, \u22120.12 eV/atom) show positive or near-zero energies, suggesting poor local stability. These values should not be interpreted as formation energies or convex-hull distances. Because CHGNet was primarily trained on inorganic crystalline materials, its predictions for molecular hybrid crystals should be interpreted as approximate single-point rankings. Interestingly, the DFT-assessed metallic structures (0742, 0789, 0391) have moderate CHGNet energies (\u22124.72 to \u22123.41 eV/atom), indicating they are locally stable but may not represent global minima. No clear correlation is observed between CHGNet energy and band-gap proxy (Pearson r \u2248 0.1), confirming that relative energy and electronic properties are independent screening criteria.")

# ===== FIGURE 3: Diversity =====
add_image(doc, 
    os.path.join(BASE_DIR, "figures", "figure_3_diversity.png"),
    width=Inches(5.5),
    caption="Fig. 3 | Chemical diversity of charge-neutral halide-containing hybrid structures. a, Metal center distribution across 50 halide structures, with Ni (7), Mn (6) and Be (5) being most frequent. b, Halide anion occurrence distribution, with Br (18 occurrences) and I (15) most common; percentages sum to slightly above 100% because one structure contains both Br and Cl. c, Organic template class distribution within the halide subset; TriA (X210), PentaHexaA (X212), and DiA-alkyl (X208) are the most frequent classes. d, Metal–halide co-occurrence heatmap showing structural diversity across the periodic table. e, Organic class versus halide co-occurrence heatmap, showing the preferential pairing of higher-charge and rod-like templates with bromides and iodides.")

add_paragraph(doc, "Table 3. Summary statistics for 50 halide-containing structures.", bold=True)
headers = ["Property", "Mean", "Median", "Std Dev", "Min", "Max"]
rows = [
    ["Band-gap proxy (eV)", "1.26", "1.20", "0.82", "0.2", "3.8"],
    ["CHGNet energy (eV/atom)", "\u22122.52", "\u22122.93", "1.77", "\u22124.99", "+1.67"],
    ["Number of atoms", "29.6", "33", "13.5", "11", "69"],
]
add_table_from_data(doc, headers, rows)

add_paragraph(doc, "Table 4. Representative halide structures across energy and band-gap proxy ranges.", bold=True)
headers = ["Crystal", "Formula", "Halide", "Metal", "Atoms", "Band-gap proxy (eV)", "CHGNet (eV/atom)", "Energy ranking"]
rows = [
    ["0217", "MnH$_{8}$C$_{3}$I$_{3}$N", "I", "Mn", "16", "0.5", "\u22124.99", "Lowest energy"],
    ["0672", "FeH$_{8}$C$_{3}$Br$_{3}$N", "Br", "Fe", "16", "0.7", "\u22124.98", "Lowest energy"],
    ["0631", "CsH$_{6}$CBr$_{2}$N", "Br", "Cs", "11", "1.7", "\u22124.42", "Favorable"],
    ["0396", "H$_{36}$C$_{7}$N$_{13}$F$_{13}$", "F", "\u2014", "69", "3.8", "\u22123.69", "Favorable"],
    ["0048", "La$_{2}$H$_{30}$C$_{5}$Br$_{11}$N$_{5}$", "Br", "La", "53", "1.4", "\u22120.12", "Near-zero energy"],
    ["0248", "Ni$_{2}$H$_{18}$C$_{3}$I$_{7}$N$_{3}$", "I", "Ni", "33", "0.2", "+1.67", "Positive energy"],
    ["0283", "CoNiH$_{18}$C$_{3}$I$_{7}$N$_{3}$", "I", "Co,Ni", "33", "0.2", "+0.44", "Positive energy"],
]
add_table_from_data(doc, headers, rows)
add_paragraph(doc, "The full list of 50 halide structures with complete property data is provided in SI Table S14.")

# ===== FIGURE 4: Screening =====
add_image(doc, 
    os.path.join(BASE_DIR, "figures", "figure_4_screening.png"),
    width=Inches(5.5),
    caption="Fig. 4 | Preliminary screening landscape of halide-containing candidates. a, CHGNet single-point energy versus band-gap proxy for 50 halide structures, colour-coded by halide type and sized by DFT selection. Low-energy candidates include crystal_0217, crystal_0672, crystal_0912 and crystal_0631, whereas positive-energy candidates include crystal_0248 and crystal_0283. CHGNet energies are single-point predictions and are not formation energies. b, CHGNet energy distribution, spanning from −4.99 to +1.67 eV/atom (median = −3.0 eV/atom). c, Distribution of band-gap proxy estimates, showing bimodal behaviour with peaks at narrow gaps (~0.2–0.5 eV, transition-metal halides) and wide gaps (>2 eV, alkali/alkaline-earth halides and organic fluorides).")

add_heading(doc, "3.4 DFT Single-Point Assessment", level=2)
add_paragraph(doc, "We selected ten representative structures for DFT-PBE single-point calculations to assess electronic structure across diverse chemistries (Table 5). All calculations were non-spin-polarized; no geometry relaxation was performed. The reported band gaps represent single-point electronic-structure assessments of the as-generated geometries.")

add_paragraph(doc, "Table 5. DFT-PBE single-point results for ten representative structures.", bold=True)
headers = ["Crystal", "Formula", "Atoms", "Metal", "Halide", "SCF iter.", "Total E (Ry)", "Fermi (eV)", "Indirect gap (eV)", "Direct gap (eV)", "Type", "Direct/Indirect"]
rows = [
    ["0742", "NiH$_{8}$C$_{3}$I$_{3}$N", "16", "Ni", "I", "15", "\u22121564.833", "1.880", "0.00", "0.06", "Metallic", "\u2014"],
    ["0789", "NiH$_{8}$C$_{3}$I$_{3}$N", "16", "Ni", "I", "13", "\u22121564.848", "3.416", "0.00", "0.03", "Metallic", "\u2014"],
    ["0391", "Mg$_{2}$H$_{12}$C$_{2}$I$_{6}$N$_{2}$", "24", "Mg", "I", "14", "\u22122430.563", "1.856", "0.00", "0.07", "Metallic", "\u2014"],
    ["0631", "CsH$_{6}$CBr$_{2}$N", "11", "Cs", "Br", "8", "\u2212202.726", "0.525", "2.17", "2.49", "Semiconductor", "Indirect"],
    ["0912", "KH$_{8}$C$_{3}$NF$_{2}$", "15", "K", "F", "12", "\u2212293.642", "4.268", "3.27", "3.39", "Semiconductor", "Indirect"],
    ["0217", "MnH$_{8}$C$_{3}$I$_{3}$N", "16", "Mn", "I", "14", "\u22121432.582", "4.668", "0.46", "0.49", "Semiconductor", "Indirect"],
    ["0672", "FeH$_{8}$C$_{3}$Br$_{3}$N", "16", "Fe", "Br", "16", "\u2212554.424", "4.038", "0.54", "0.60", "Semiconductor", "Indirect"],
    ["0927", "CoH$_{12}$C$_{2}$N$_{2}$Cl$_{4}$", "21", "Co", "Cl", "12", "\u2212516.600", "2.244", "0.54", "0.56", "Semiconductor", "Indirect"],
    ["0626", "VH$_{11}$C$_{2}$Br$_{6}$N$_{3}$", "23", "V", "Br", "32", "\u2212534.121", "\u22122.633", "0.78", "0.78", "Semiconductor", "Direct"],
    ["0632", "CrH$_{11}$C$_{2}$Br$_{6}$N$_{3}$", "23", "Cr", "Br", "36", "\u2212565.254", "\u22122.017", "0.51", "0.51", "Semiconductor", "Indirect"],
]
add_table_from_data(doc, headers, rows)

add_paragraph(doc, "Three structures (0742, 0789, 0391) exhibit zero band gaps and metallic character. The isomer pair 0742 and 0789 (same formula NiH$_{8}$C$_{3}$I$_{3}$N, different lattice parameters) both show metallic behavior, with 0789 being more stable by 0.015 Ry. The Mg-based structure (0391) also shows a zero gap; we note it has a non-physical Mg\u2013C distance of 1.27 \u00c5 (vs. expected 2.2\u20132.5 \u00c5), indicating possible atomic overlap that may affect the electronic structure.", indent=True)
add_paragraph(doc, "Crystal_0631 (CsH$_{6}$CBr$_{2}$N) shows a wide indirect band gap of 2.17 eV, with a direct gap of 2.49 eV at the \u0393 point, making it a candidate wide-band-gap semiconductor. Crystal_0912 (KH$_{8}$C$_{3}$NF$_{2}$) exhibits an even larger indirect gap of 3.27 eV, consistent with the high ionicity of the K\u2013F bond. Among the transition-metal halides, crystal_0626 (VH$_{11}$C$_{2}$Br$_{6}$N$_{3}$) is the **only** structure with a direct band gap (0.78 eV), which shows potentially interesting direct-gap behavior. The remaining semiconductor structures (0217, 0672, 0927, 0632) all exhibit indirect gaps in the 0.46\u20130.78 eV range, typical of transition-metal halides where d-electron states contribute to the band edges.", indent=True)
add_paragraph(doc, "The diversity of electronic behavior\u2014ranging from metallic (Ni, Mg iodides) to wide-gap semiconducting (Cs bromide, K fluoride) to narrow-gap direct-band-gap (V bromide)\u2014suggests that the virtual-element framework can generate chemically diverse candidates with plausible electronic behavior, while quantitative validation requires relaxation and higher-level calculations. However, because no geometry relaxation was performed, the true equilibrium gaps may differ from the reported single-point values. Furthermore, the non-spin-polarized approximation may underestimate gaps for magnetic systems, although this is less relevant for the closed-shell Cs- and K-based systems. Full DFT computational details, input files, and convergence diagnostics are provided in SI Section S2.", indent=True)

# ===== FIGURE 5: DFT Assessment =====
add_image(doc, 
    os.path.join(BASE_DIR, "figures", "figure_5_dft.png"),
    width=Inches(5.5),
    caption="Fig. 5 | DFT-PBE single-point assessment of ten representative structures. a, DFT-PBE indirect band gap versus CHGNet energy for the ten selected candidates, colour-coded by band-gap type (metallic, indirect, direct). Metallic structures with negative indirect gaps are plotted at zero gap for visualization. b, Direct versus indirect band-gap comparison for the seven semiconducting structures, showing that crystal_0626 (V bromide) is the only direct-gap candidate (0.78 eV), while all others are indirect. c, Total density of states (DOS) for crystal_0631 (CsH$_{6}$CBr$_{2}$N), with the Fermi level indicated by a dashed line. d, Band-energy distribution of crystal_0631 over the sampled 4×4×4 k-point grid (112 k-points). All DFT calculations are non-spin-polarized single-point PBE calculations on as-generated structures.")

# Section 4
add_heading(doc, "4. Discussion", level=1)

add_heading(doc, "4.1 Why Virtual Elements Are Effective", level=2)
add_paragraph(doc, "The virtual element framework successfully compresses the vast organic cation space into 12 manageable classes. Organic cations in hybrid crystals primarily provide: (i) charge compensation, (ii) spatial templating, (iii) hydrogen bonding or electrostatic interactions, and (iv) dimensionality control. By capturing these first-order crystal-chemical roles through coarse-grained classes, the virtual-element approach preserves the essential structural information needed for crystal generation while remaining compatible with atomic-type diffusion models. This is analogous to coarse-grained molecular dynamics, where detailed atomistic degrees of freedom are replaced by effective interaction sites.", indent=True)
add_paragraph(doc, "However, intra-class diversity is lost during generation: all A3D-small cations (MA, FA, GA, AZ) map to the same virtual element X201. Future work could employ continuous embeddings to preserve finer distinctions (SI Section S5).", indent=True)

add_heading(doc, "4.2 Charge Balance as a Missing Inductive Bias", level=2)
add_paragraph(doc, "The low charge-neutral rate (11.1%) reflects a fundamental challenge in generative materials models: the diffusion process optimizes local structural features through a per-atom denoising objective without global charge constraints. Primary sources of imbalance include insufficient cationic charge compensation, oxidation-state ambiguity for multi-valent metals, and organic cation overcounting. This is not a failure of the model per se, but a revelation that current crystal diffusion models lack charge conservation as an inductive bias.", indent=True)
add_paragraph(doc, "Potential solutions include adding a differentiable charge-balance penalty during training, constraining metal:halide ratios during generation, or employing charge-conditioned diffusion sampling. These directions are promising for improving the practical utility of generative hybrid crystal models.", indent=True)

add_heading(doc, "4.3 DFT Assessment Insights and Limitations", level=2)
add_paragraph(doc, "The DFT results reveal several important insights. First, the model generates both metallic and semiconducting structures, demonstrating diversity in electronic properties. Second, the zero-gap results for transition-metal systems highlight the need for post-generation geometry optimization (vc-relax) and beyond-PBE methods (PBE+U, HSE06) for accurate band-gap prediction. The alkali-metal halide structures (crystal_0631 with 2.17 eV indirect gap and crystal_0912 with 3.27 eV indirect gap) are particularly promising as they avoid d-electron correlation complexities, but their true stability and electronic properties require further validation after structural relaxation.", indent=True)
add_paragraph(doc, "We emphasize several limitations of the current DFT assessment: (1) no geometry relaxation was performed, so the true equilibrium geometries may differ significantly from the as-generated structures; (2) all calculations were non-spin-polarized, so magnetic ordering in transition-metal systems is not captured; (3) PBE often underestimates band gaps, but geometry relaxation and magnetic ordering may shift band edges in either direction; the reported values should therefore be regarded as preliminary single-point estimates; (4) the indirect/direct gap classification is based on a 4\u00d74\u00d74 k-point grid, which may not fully capture the true band dispersion; (5) only one structure (0626) exhibits a direct band gap, and the narrow indirect gaps of the transition-metal halides may not be suitable for photovoltaic applications without further band-structure engineering. These limitations are standard for preliminary computational screening and do not diminish the value of the virtual-element framework, but they highlight the need for follow-up validation studies on the most promising candidates.", indent=True)

add_heading(doc, "4.4 Comparison to Prior Work and Broader Applicability", level=2)
add_paragraph(doc, "Previous machine-learning studies on hybrid perovskites have primarily focused on property prediction, synthesis classification, compositional recommendation or inverse design of organic cation descriptors [11\u201317]. These approaches are valuable for navigating known or enumerated chemical spaces but generally do not generate full three-dimensional hybrid crystal structures containing molecular organic cations. Existing methods for hybrid perovskite design also include crystal structure prediction (CSP) via evolutionary algorithms [6,7], machine learning potentials for structure relaxation [8], and generative models for inorganic composition generation [9,10]. Our approach differs by enabling end-to-end structure generation for hybrid materials using a pretrained inorganic diffusion model, rather than developing a new model from scratch. While CSP methods achieve higher charge balance rates (typically >50%) through explicit energy minimization, they are computationally expensive (weeks vs. hours). A hybrid approach\u2014generative sampling followed by CSP relaxation\u2014may offer the optimal balance of efficiency and accuracy.", indent=True)
add_paragraph(doc, "The virtual-element framework is not limited to hybrid perovskites. It can be extended to layered organic\u2013inorganic frameworks, hybrid halides, metal\u2013organic frameworks (MOFs), organic-templated zeolites, and molecular salts, provided that the organic component can be classified into coarse-grained functional classes.", indent=True)

# Section 5
add_heading(doc, "5. Conclusion", level=1)
add_paragraph(doc, "We have introduced a virtual-element framework that extends a generative diffusion model for inorganic materials to organic\u2013inorganic hybrid crystals. The central contribution is a representation layer that makes molecular organic ions compatible with atom-type crystal diffusion models, enabling pretrained inorganic generative models to enter hybrid crystal space without retraining from scratch. By mapping organic cation families to 12 active virtual-element classes, we expanded the atom-type index space of MatterGen while preserving pretrained inorganic embeddings, and we developed a molecular-template decoding workflow that converts coarse-grained virtual crystals into full-atom hybrid structures. From 1,000 generated structures, 954 were successfully decoded and 106 were charge-neutral, of which 50 contained halides and spanned diverse metal centers and anion chemistries. DFT-PBE single-point assessment of ten representative candidates revealed a spectrum of electronic behavior, from metallic transition-metal iodides to a direct-gap V-based bromide and wide-gap indirect-gap Cs\u2013Br and K\u2013F systems.", indent=True)
add_paragraph(doc, "Beyond the specific candidates identified here, this work establishes a transferable representation, generation and decoding framework that broadens the scope of pretrained inorganic crystal generators to the vast, compositionally flexible space of organic\u2013inorganic hybrid materials. It also highlights charge neutrality as a critical missing inductive bias in current crystal diffusion models and provides a practical baseline for future charge-constrained generative approaches. By demonstrating that molecular organic components can be integrated into atom-type diffusion models through coarse-grained tokens, this framework opens a route toward computational exploration of hybrid halides, layered frameworks, molecular salts and related materials that lie beyond the reach of element-limited generative models.", indent=True)

# Data and Code Availability
add_heading(doc, "Data and Code Availability", level=1)
add_paragraph(doc, "All generated structures, analysis scripts, and DFT input/output files are available in the paper_results/ directory. The training dataset was constructed from publicly available hybrid crystal structures, including perovskite-related and other framework templates, together with synthetically generated combinations. The pretrained MatterGen weights are available through Microsoft Research. Modified code including ExtendedAtomEmbedding, virtual element definitions, and decoding scripts is available at https://github.com/yqzhan-blip/virtual-element-hybrid-crystals and archived at [Zenodo DOI].")

# Acknowledgments
add_heading(doc, "Acknowledgments", level=1)
add_paragraph(doc, "The authors acknowledge funding support from the Shanghai Science and Technology Innovation Action Plan (No. 24DZ3001200), the National Natural Science Foundation of China (No. 62304046 and No. 62274040), and the Basic Research Project of State Key Laboratory of Photovoltaic Science and Technology (No. 202401020302).")

# Author Contributions
add_heading(doc, "Author Contributions", level=1)
add_paragraph(doc, "Y.Z. conceived the project, designed the virtual-element framework, implemented the model extensions and decoding pipeline, performed generation and DFT validation, analyzed the data, and wrote the manuscript.")

# Competing Interests
add_heading(doc, "Competing Interests", level=1)
add_paragraph(doc, "The authors declare no competing interests.")

# References
add_heading(doc, "References", level=1)
refs = [
    "[1] Kojima, A., Teshima, K., Shirai, Y. & Miyasaka, T. Organometal halide perovskites as visible-light sensitizers for photovoltaic cells. *J. Am. Chem. Soc.* **131**, 6050\u20136051 (2009).",
    "[2] Green, M. A., Ho-Baillie, A. & Snaith, H. J. The emergence of perovskite solar cells. *Nat. Photon.* **8**, 506\u2013514 (2014).",
    "[3] Zeni, C., et al. MatterGen: a generative model for inorganic materials design. *Nature* **639**, 624\u2013632 (2025).",
    "[4] Giannozzi, P., et al. Quantum ESPRESSO: a modular and open-source software project for quantum simulations of materials. *J. Phys.: Condens. Matter* **21**, 395502 (2009).",
    "[5] Deng, B., et al. CHGNet as a pretrained universal neural network potential for charge-informed atomistic modelling. *Nat. Mach. Intell.* **5**, 1031\u20131041 (2023).",
    "[6] Oganov, A. R., et al. Crystal structure prediction using ab initio evolutionary techniques. *J. Chem. Phys.* **124**, 244704 (2006).",
    "[7] Wang, Y., et al. Crystal structure prediction via particle-swarm optimization. *Phys. Rev. B* **82**, 094116 (2010).",
    "[8] Batzner, S., et al. E(3)-equivariant graph neural networks for data-efficient and accurate interatomic potentials. *Nat. Commun.* **13**, 2453 (2022).",
    "[9] Dan, Y., et al. Generative adversarial networks (GAN) based efficient sampling of chemical composition space for inverse design of inorganic materials. *npj Comput. Mater.* **6**, 84 (2020).",
    "[10] Court, C. J., Yildirim, B., Jain, A. & Cole, J. M. 3-D inorganic crystal structure generation and property prediction via representation learning. *J. Chem. Inf. Model.* **60**, 4518\u20134535 (2020).",
    "[11] Lu, S., et al. Accelerated discovery of stable lead-free hybrid organic\u2013inorganic perovskites via machine learning. *Nat. Commun.* **9**, 3405 (2018).",
    "[12] Lyu, R., et al. Predictive design model for low-dimensional organic\u2013inorganic halide perovskites assisted by machine learning. *J. Am. Chem. Soc.* **143**, 12766\u201312776 (2021).",
    "[13] Wu, Y., et al. Universal machine learning aided synthesis approach of two-dimensional perovskites in a typical laboratory. *Nat. Commun.* **15**, 14 (2024).",
    "[14] Sun, J. & Wu, Y. Organic\u2013inorganic metal halide perovskites: toward stability, chirality, and AI-guided discovery. *ACS Cent. Sci.* **12**, 399\u2013412 (2026).",
    "[15] Ai, Q., Norquist, A. J. & Schrier, J. Predicting compositional changes of organic\u2013inorganic hybrid materials with Augmented CycleGAN. *Digit. Discov.* **1**, 255\u2013265 (2022).",
    "[16] Husain, M. S. Perspective\u2014accelerated discovery of organic-inorganic hybrid materials via machine learning. *ECS J. Solid State Sci. Technol.* **10**, 037001 (2021).",
    "[17] Lyu, Y., et al. Fingerprinting organic molecules for the inverse design of two-dimensional hybrid perovskites with target energetics. *Sci. Adv.* **12**, eaeb4144 (2026).",
]
for ref in refs:
    p = add_paragraph(doc, ref)
    for run in p.runs:
        run.font.size = Pt(10)

# Supplementary Information reference
add_heading(doc, "Supplementary Information", level=1)
add_paragraph(doc, "The following sections are available in the accompanying Supplementary Information document:")
add_paragraph(doc, "\u2022 S1. Dataset construction, training hyperparameters, and 3D template library details")
add_paragraph(doc, "\u2022 S2. Complete DFT-PBE single-point assessment for ten structures: computational details, input files, convergence diagnostics, full energy tables, band structure analysis, and electronic structure figures")
add_paragraph(doc, "\u2022 S3. Complete charge balance analysis with full structure list (106 entries)")
add_paragraph(doc, "\u2022 S4. Complete halide structure list (50 entries) with band-gap proxies and CHGNet energy data")
add_paragraph(doc, "\u2022 S5. Continuous embedding strategy for finer organic cation differentiation")
add_paragraph(doc, "\u2022 S6. Full virtual element distribution and template selection statistics")

# Save
output_path = os.path.join(BASE_DIR, "paper_draft_v5.docx")
doc.save(output_path)
print(f"Saved: {output_path}")

# Verify
print("\nFigure insertion verification:")
fig_names = ["figure_1_framework.png", "figure_2_statistics.png", 
             "figure_3_diversity.png", "figure_4_screening.png",
             "figure_5_dft.png"]
for fn in fig_names:
    p = os.path.join(BASE_DIR, "figures", fn)
    print(f"  {'OK' if os.path.exists(p) else 'MISSING'}: {fn}")
print("\nDone!")
