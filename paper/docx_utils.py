"""Shared utilities for generating DOCX documents with markdown/math formatting."""
import re
from docx.shared import Pt
from docx.oxml.ns import qn


def set_chinese_font(run, font_name="Times New Roman", size=11, bold=False, italic=False, color=None):
    """Apply font formatting to a run, including optional font color."""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color is not None:
        run.font.color.rgb = color


def tokenize_math(math_text):
    """Tokenize LaTeX-like math into (fmt, content) tuples.

    fmt: 'normal', 'sub', 'super'
    content: str or list of nested tokens
    """
    tokens = []
    i = 0
    n = len(math_text)
    buf = ''
    while i < n:
        if i < n - 1 and math_text[i] in '_^':
            if math_text[i + 1] == '{':
                if buf:
                    tokens.append(('normal', buf))
                    buf = ''
                depth = 1
                j = i + 2
                inner = ''
                while j < n and depth > 0:
                    if math_text[j] == '{':
                        depth += 1
                        inner += '{'
                    elif math_text[j] == '}':
                        depth -= 1
                        if depth > 0:
                            inner += '}'
                    else:
                        inner += math_text[j]
                    j += 1
                fmt = 'sub' if math_text[i] == '_' else 'super'
                tokens.append((fmt, tokenize_math(inner)))
                i = j
            else:
                # Single-character subscript/superscript (e.g., _2 or ^+)
                if buf:
                    tokens.append(('normal', buf))
                    buf = ''
                fmt = 'sub' if math_text[i] == '_' else 'super'
                tokens.append((fmt, math_text[i + 1]))
                i += 2
        else:
            buf += math_text[i]
            i += 1
    if buf:
        tokens.append(('normal', buf))
    return tokens


def add_math_runs(paragraph, math_text, base_bold=False, base_italic=True, size=11):
    """Add math content as italic runs with proper subscripts/superscripts."""
    tokens = tokenize_math(math_text)
    _add_math_tokens(paragraph, tokens, base_bold, base_italic, size)


def _add_math_tokens(paragraph, tokens, base_bold, base_italic, size, inherited_fmt='normal'):
    for fmt, content in tokens:
        # The effective format is the explicit token format if it is sub/super,
        # otherwise inherit from the parent (needed for nested tokenization).
        effective_fmt = fmt if fmt in ('sub', 'super') else inherited_fmt
        if isinstance(content, list):
            _add_math_tokens(paragraph, content, base_bold, base_italic, size, effective_fmt)
        else:
            run = paragraph.add_run(content)
            set_chinese_font(run, bold=base_bold, italic=base_italic, size=size)
            if effective_fmt == 'sub':
                run.font.subscript = True
            elif effective_fmt == 'super':
                run.font.superscript = True


def add_formatted_runs(paragraph, text, base_bold=False, base_italic=False, size=11):
    """Parse **bold**, *italic*, and $math$ in text and add runs with proper formatting.

    Inside $...$ math expressions, _{...} is rendered as subscript and ^{...} as superscript.
    """
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|\$.*?\$)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run_text = part[2:-2]
            if run_text:
                run = paragraph.add_run(run_text)
                set_chinese_font(run, bold=True, italic=base_italic, size=size)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run_text = part[1:-1]
            if run_text:
                run = paragraph.add_run(run_text)
                set_chinese_font(run, bold=base_bold, italic=True, size=size)
        elif part.startswith('$') and part.endswith('$'):
            math_content = part[1:-1]
            if math_content:
                add_math_runs(paragraph, math_content, base_bold=base_bold,
                              base_italic=True, size=size)
        else:
            run = paragraph.add_run(part)
            set_chinese_font(run, bold=base_bold, italic=base_italic, size=size)
