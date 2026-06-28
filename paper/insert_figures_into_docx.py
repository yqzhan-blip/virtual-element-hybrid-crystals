"""Insert Figure 1-5 into paper_draft_v5.docx at correct positions."""
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

def set_chinese_font(run, font_name="Times New Roman", size=11, bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_image_with_caption(doc, image_path, width=Inches(5.5), caption=""):
    if not os.path.exists(image_path):
        print(f"  [SKIP] Image not found: {image_path}")
        return False
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=width)
    
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(caption)
        set_chinese_font(cr, italic=True, size=10)
    
    doc.add_paragraph()  # spacing
    print(f"  [OK] Inserted: {os.path.basename(image_path)}")
    return True

# Load existing document
doc_path = os.path.join(BASE_DIR, "paper_draft_v5.docx")
doc = Document(doc_path)

print(f"Loaded document: {doc_path}")
print(f"Total paragraphs: {len(doc.paragraphs)}")

# Figure definitions
fig_dir = os.path.join(BASE_DIR, "figures")
analysis_dir = os.path.join(BASE_DIR, "analysis", "charts")

figures = [
    {
        "search_text": "Generation and Decoding Pipeline",
        "image": os.path.join(fig_dir, "figure_1_pipeline.png"),
        "caption": "Figure 1. Overview of the virtual-element generation and decoding pipeline. (a) The fine-tuned MatterGen model generates coarse-grained structures with virtual elements (X201–X212). (b) The decoding pipeline maps each virtual site to a full-atom organic cation using a precomputed 3D molecular template library.",
        "after_table": True,
        "heading_level": 2,
    },
    {
        "search_text": "Charge Balance Analysis",
        "image": os.path.join(fig_dir, "figure_2_charge_distribution.png"),
        "caption": "Figure 2. Net charge distribution of 954 decoded structures. The majority (893, 93.6%) are unbalanced, with the most common net charges being Q = −5 (128), −1 (158), and −3 (110). Only 106 structures (11.1%) satisfy charge neutrality.",
        "after_table": True,
        "heading_level": 2,
    },
    {
        "search_text": "Halide-containing Structures",
        "image": os.path.join(fig_dir, "figure_3_element_statistics.png"),
        "caption": "Figure 3. Element and template distribution in 50 halide-containing charge-balanced structures. (a) Metal center distribution. (b) Halide anion distribution. (c) Organic template type distribution.",
        "after_table": False,
        "heading_level": 2,
    },
    {
        "search_text": "Representative halide structures",
        "image": os.path.join(fig_dir, "figure_4_representative_structures.png"),
        "caption": "Figure 4. Representative crystal structures from the halide subset. (a) crystal_0217 (MnH₈C₃I₃N, lowest CHGNet energy). (b) crystal_0631 (CsH₆CBr₂N, wide band gap). (c) crystal_0912 (KH₈C₃NF₂, highest band gap).",
        "after_table": True,
        "heading_level": None,  # special: search in paragraph text
        "search_in_para": True,
    },
    {
        "search_text": "Table 3. Summary statistics",
        "image": os.path.join(analysis_dir, "bandgap_histogram.png"),
        "caption": "Figure 5. Distribution of band-gap proxies for 50 halide-containing structures. The histogram shows a bimodal distribution with peaks at narrow gaps (~0.2–0.5 eV, transition-metal halides) and wide gaps (>2 eV, alkali/alkaline-earth halides and organic fluorides).",
        "after_table": True,
        "heading_level": None,
        "search_in_para": True,
    },
]

inserted = 0
for fig_info in figures:
    found = False
    
    # Determine search strategy
    if fig_info.get("search_in_para"):
        # Search in paragraph text
        for i, para in enumerate(doc.paragraphs):
            if fig_info["search_text"] in para.text:
                # Insert after this paragraph
                insert_idx = i + 1
                # Build a temporary paragraph at the end and move it
                # Actually easier: insert a new paragraph after current
                # Use the element approach
                p = doc.add_paragraph()
                p._element.getparent().remove(p._element)
                para._element.addnext(p._element)
                
                # Add image to this new paragraph
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                if os.path.exists(fig_info["image"]):
                    run.add_picture(fig_info["image"], width=Inches(5.5))
                    print(f"  [OK] Inserted Figure after paragraph: '{para.text[:60]}...'")
                else:
                    print(f"  [SKIP] Image not found: {fig_info['image']}")
                
                # Add caption
                if fig_info["caption"]:
                    cp = doc.add_paragraph()
                    cp._element.getparent().remove(cp._element)
                    p._element.addnext(cp._element)
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cr = cp.add_run(fig_info["caption"])
                    set_chinese_font(cr, italic=True, size=10)
                
                # Add spacing
                sp = doc.add_paragraph()
                sp._element.getparent().remove(sp._element)
                cp._element.addnext(sp._element)
                
                found = True
                inserted += 1
                break
    else:
        # Search in heading
        for i, para in enumerate(doc.paragraphs):
            if para.style.name.startswith("Heading") and fig_info["search_text"] in para.text:
                # Find the next table after this heading
                insert_para = para
                for j in range(i+1, min(i+20, len(doc.paragraphs))):
                    if fig_info.get("after_table") and doc.paragraphs[j]._element.getnext() is not None:
                        # Check if next element is a table
                        next_elem = doc.paragraphs[j]._element.getnext()
                        if next_elem is not None and next_elem.tag.endswith('tbl'):
                            insert_para = doc.paragraphs[j]
                            # After the table
                            tbl_next = next_elem.getnext()
                            if tbl_next is not None:
                                for k in range(j+1, len(doc.paragraphs)):
                                    if doc.paragraphs[k]._element == tbl_next:
                                        insert_para = doc.paragraphs[k]
                                        break
                            break
                
                # Insert after insert_para
                p = doc.add_paragraph()
                p._element.getparent().remove(p._element)
                insert_para._element.addnext(p._element)
                
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                if os.path.exists(fig_info["image"]):
                    run.add_picture(fig_info["image"], width=Inches(5.5))
                    print(f"  [OK] Inserted Figure after heading/paragraph: '{insert_para.text[:60]}...'")
                else:
                    print(f"  [SKIP] Image not found: {fig_info['image']}")
                
                if fig_info["caption"]:
                    cp = doc.add_paragraph()
                    cp._element.getparent().remove(cp._element)
                    p._element.addnext(cp._element)
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cr = cp.add_run(fig_info["caption"])
                    set_chinese_font(cr, italic=True, size=10)
                
                sp = doc.add_paragraph()
                sp._element.getparent().remove(sp._element)
                cp._element.addnext(sp._element)
                
                found = True
                inserted += 1
                break
    
    if not found:
        print(f"  [NOT FOUND] Could not find insertion point for: {fig_info['search_text']}")

# Save
output_path = os.path.join(BASE_DIR, "paper_draft_v5.docx")
doc.save(output_path)
print(f"\nSaved: {output_path}")
print(f"Total figures inserted: {inserted}/5")
