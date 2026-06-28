"""Generate cover_letter.docx from cover_letter.md."""
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

with open(os.path.join(BASE_DIR, 'cover_letter.md'), 'r', encoding='utf-8') as f:
    lines = f.readlines()

for line in lines:
    line = line.rstrip('\n')
    if not line.strip():
        doc.add_paragraph()
        continue
    if line.startswith('**') and line.endswith('**'):
        # Bold title
        p = doc.add_paragraph()
        run = p.add_run(line.strip('*'))
        run.bold = True
        run.font.size = Pt(12)
    else:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(6)

output_path = os.path.join(BASE_DIR, 'cover_letter.docx')
doc.save(output_path)
print(f"Saved: {output_path}")
